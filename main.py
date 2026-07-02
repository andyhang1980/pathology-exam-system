"""
全国病理学副高级职称考试模拟系统
Pathology Associate Senior Professional Title Exam Simulation System
"""

import os
import json
import re
import sqlite3
import httpx
import fitz  # PyMuPDF for PDF
from docx import Document as DocxDocument
import secrets
import hashlib
import uuid
import shutil
from datetime import datetime, timedelta
from typing import Optional
from contextlib import contextmanager
from io import BytesIO

from fastapi import FastAPI, Request, Response, Form, UploadFile, File, HTTPException, Depends
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from itsdangerous import URLSafeTimedSerializer

# ============================================================
# Password hashing (SHA-256 + salt)
# ============================================================
def hash_password(password: str) -> str:
    salt = secrets.token_hex(16)
    h = hashlib.sha256((salt + password).encode()).hexdigest()
    return f"{salt}${h}"

def verify_password(password: str, hashed: str) -> bool:
    try:
        salt, h = hashed.split("$", 1)
        return hashlib.sha256((salt + password).encode()).hexdigest() == h
    except (ValueError, AttributeError):
        return False

# ============================================================
# Configuration
# ============================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DATA_DIR = os.path.join(BASE_DIR, "data")
DB_PATH = os.path.join(DATA_DIR, "exam.db")
SECRET_KEY = os.environ.get("SECRET_KEY", secrets.token_hex(32))

os.makedirs(DATA_DIR, exist_ok=True)

UPLOAD_DIR = os.path.join(BASE_DIR, "static", "uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ============================================================
# FastAPI App
# ============================================================
app = FastAPI(title="张老师病理学职称考试模拟系统")
app.mount("/static", StaticFiles(directory=os.path.join(BASE_DIR, "static")), name="static")
templates = Jinja2Templates(directory=os.path.join(BASE_DIR, "templates"))
serializer = URLSafeTimedSerializer(SECRET_KEY)

# ============================================================
# Database
# ============================================================
@contextmanager
def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    conn.execute("PRAGMA journal_mode=WAL")
    conn.execute("PRAGMA foreign_keys=ON")
    try:
        yield conn
        conn.commit()
    except Exception:
        conn.rollback()
        raise
    finally:
        conn.close()

def init_db():
    with get_db() as db:
        db.executescript("""
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            username TEXT UNIQUE NOT NULL,
            password TEXT NOT NULL,
            real_name TEXT DEFAULT '',
            student_id TEXT DEFAULT '',
            login_type TEXT DEFAULT 'account',
            role TEXT DEFAULT 'student',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            type TEXT NOT NULL CHECK(type IN ('single','multiple','shared','case')),
            category TEXT DEFAULT '',
            content TEXT NOT NULL,
            options TEXT DEFAULT '[]',
            answer TEXT NOT NULL,
            explanation TEXT DEFAULT '',
            score REAL DEFAULT 1,
            sub_questions TEXT DEFAULT '[]',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS exams (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            description TEXT DEFAULT '',
            duration INTEGER DEFAULT 120,
            total_score REAL DEFAULT 100,
            pass_score REAL DEFAULT 60,
            is_practice BOOLEAN DEFAULT 0,
            shuffle_questions BOOLEAN DEFAULT 0,
            status TEXT DEFAULT 'active',
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        );

        CREATE TABLE IF NOT EXISTS exam_questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            exam_id INTEGER NOT NULL,
            question_id INTEGER NOT NULL,
            display_order INTEGER DEFAULT 0,
            score_override REAL,
            FOREIGN KEY (exam_id) REFERENCES exams(id) ON DELETE CASCADE,
            FOREIGN KEY (question_id) REFERENCES questions(id) ON DELETE CASCADE
        );

        CREATE TABLE IF NOT EXISTS exam_sessions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER NOT NULL,
            exam_id INTEGER NOT NULL,
            start_time TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            end_time TIMESTAMP,
            total_score REAL DEFAULT 0,
            status TEXT DEFAULT 'in_progress',
            FOREIGN KEY (user_id) REFERENCES users(id),
            FOREIGN KEY (exam_id) REFERENCES exams(id)
        );

        CREATE TABLE IF NOT EXISTS session_answers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            session_id INTEGER NOT NULL,
            question_id INTEGER NOT NULL,
            sub_index INTEGER DEFAULT 0,
            answer TEXT DEFAULT '',
            is_correct BOOLEAN DEFAULT 0,
            score REAL DEFAULT 0,
            FOREIGN KEY (session_id) REFERENCES exam_sessions(id) ON DELETE CASCADE,
            FOREIGN KEY (question_id) REFERENCES questions(id) ON DELETE CASCADE
        );
        """)

        db.execute("""CREATE TABLE IF NOT EXISTS textbooks (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            file_name TEXT NOT NULL,
            file_size INTEGER DEFAULT 0,
            content TEXT NOT NULL DEFAULT '',
            char_count INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )""")
        # Migration: add images column if not exists
        cols = [row["name"] for row in db.execute("PRAGMA table_info(questions)").fetchall()]
        if "images" not in cols:
            db.execute("ALTER TABLE questions ADD COLUMN images TEXT DEFAULT '[]'")
        if "source_set" not in cols:
            db.execute("ALTER TABLE questions ADD COLUMN source_set TEXT DEFAULT ''")

        # Migration: add student_id and login_type to users
        user_cols = [row["name"] for row in db.execute("PRAGMA table_info(users)").fetchall()]
        if "student_id" not in user_cols:
            db.execute("ALTER TABLE users ADD COLUMN student_id TEXT DEFAULT ''")
        if "login_type" not in user_cols:
            db.execute("ALTER TABLE users ADD COLUMN login_type TEXT DEFAULT 'account'")

        # Create site_settings table
        db.execute("""CREATE TABLE IF NOT EXISTS site_settings (
            id INTEGER PRIMARY KEY CHECK (id = 1),
            allow_register INTEGER DEFAULT 1,
            allow_account_login INTEGER DEFAULT 1,
            allow_student_id_login INTEGER DEFAULT 1
        )""")
        db.execute('INSERT OR IGNORE INTO site_settings (id) VALUES (1)')

        # Migration: add last_question_index and last_sub_index to exam_sessions
        es_cols = [row["name"] for row in db.execute("PRAGMA table_info(exam_sessions)").fetchall()]
        if "last_question_index" not in es_cols:
            db.execute("ALTER TABLE exam_sessions ADD COLUMN last_question_index INTEGER DEFAULT 0")
        if "last_sub_index" not in es_cols:
            db.execute("ALTER TABLE exam_sessions ADD COLUMN last_sub_index INTEGER DEFAULT 0")

        # Migration: site_settings columns
        ss_cols = [row['name'] for row in db.execute('PRAGMA table_info(site_settings)').fetchall()]
        if 'allow_register' not in ss_cols:
            db.execute('ALTER TABLE site_settings ADD COLUMN allow_register INTEGER DEFAULT 1')
        if 'allow_account_login' not in ss_cols:
            db.execute('ALTER TABLE site_settings ADD COLUMN allow_account_login INTEGER DEFAULT 1')
        if 'allow_student_id_login' not in ss_cols:
            db.execute('ALTER TABLE site_settings ADD COLUMN allow_student_id_login INTEGER DEFAULT 1')
        user_cols2 = [c2['name'] for c2 in db.execute('PRAGMA table_info(users)').fetchall()]
        if 'plain_password' not in user_cols2:
            db.execute('ALTER TABLE users ADD COLUMN plain_password TEXT DEFAULT ""')


        # Create ads table
        db.execute("""CREATE TABLE IF NOT EXISTS ads (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL DEFAULT '',
            image_url TEXT NOT NULL,
            link_url TEXT DEFAULT '',
            is_active INTEGER DEFAULT 1,
            sort_order INTEGER DEFAULT 0,
            created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
        )""")

        # Migrate existing ad images from static/uploads/ads/ into the ads table
        ads_dir = os.path.join(UPLOAD_DIR, 'ads')
        if os.path.isdir(ads_dir):
            existing = set(row[0] for row in db.execute('SELECT image_url FROM ads').fetchall())
            for fname in sorted(os.listdir(ads_dir)):
                if fname.lower().endswith(('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp')):
                    url = f'/static/uploads/ads/{fname}'
                    if url not in existing:
                        title = os.path.splitext(fname)[0]
                        db.execute('INSERT INTO ads (title, image_url, is_active, sort_order) VALUES (?, ?, 1, 0)', (title, url))


        # Create default admin if not exists
        admin = db.execute("SELECT id FROM users WHERE username='admin'").fetchone()
        if not admin:
            db.execute(
                "INSERT INTO users (username, password, real_name, role) VALUES (?, ?, ?, ?)",
                ("admin", hash_password("admin123"), "系统管理员", "admin")
            )

        # Create demo student if not exists
        student = db.execute("SELECT id FROM users WHERE username='student'").fetchone()
        if not student:
            db.execute(
                "INSERT INTO users (username, password, real_name, role) VALUES (?, ?, ?, ?)",
                ("student", hash_password("student123"), "演示学生", "student")
            )

# ============================================================
# Auth helpers
# ============================================================
def get_current_user(request: Request) -> Optional[dict]:
    token = request.cookies.get("session_token")
    if not token:
        return None
    try:
        data = serializer.loads(token, max_age=86400 * 7)
        with get_db() as db:
            user = db.execute("SELECT * FROM users WHERE id=?", (data["user_id"],)).fetchone()
            if user:
                return dict(user)
    except Exception:
        pass
    return None

def require_login(request: Request):
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=307, headers={"Location": "/login"})
    return user

def require_admin(request: Request):
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=307, headers={"Location": "/login"})
    if user["role"] != "admin":
        raise HTTPException(status_code=403, detail="权限不足")
    return user

def make_token(user_id: int) -> str:
    return serializer.dumps({"user_id": user_id})

# ============================================================
# Scoring logic
# ============================================================
def score_question(q_type: str, user_answer: str, correct_answer: str, base_score: float = 1.0, option_weights: dict = None) -> tuple:
    """Returns (is_correct, score_earned)
    option_weights: {option_label: weight_value} for case type questions with custom weights per option
    """
    if q_type == "single":
        if user_answer.strip().upper() == correct_answer.strip().upper():
            return (True, base_score)
        return (False, 0)

    elif q_type == "multiple":
        user_set = set(user_answer.strip().upper().replace(",", ""))
        correct_set = set(correct_answer.strip().upper().replace(",", ""))
        if user_set == correct_set:
            return (True, base_score)
        return (False, 0)

    elif q_type == "shared":
        if user_answer.strip().upper() == correct_answer.strip().upper():
            return (True, base_score)
        return (False, 0)

    elif q_type == "case":
        # 案例分析题：每个提问1个得分点，正确答案为1个或多个选项
        # 支持每个选项不同权重（option_weights参数），无权重时等权
        # 少选：按选对的选项得对应分值
        # 多选、错选：倒扣对应选项的分值
        # 保底规则：本题最低得分为0分，不会出现负分
        try:
            user_selections = json.loads(user_answer) if user_answer else []
            correct_selections = json.loads(correct_answer) if correct_answer else []
        except (json.JSONDecodeError, TypeError):
            user_selections = [x.strip().upper() for x in user_answer.split(",") if x.strip()] if user_answer else []
            correct_selections = [x.strip().upper() for x in correct_answer.split(",") if x.strip()] if correct_answer else []

        user_set = set(user_selections)
        correct_set = set(correct_selections)

        if option_weights and len(option_weights) > 0:
            # 自定义权重方案：每个选项有不同权重
            # 计算正确选项的总权重，归一化到 base_score
            total_correct_weight = sum(option_weights.get(sel, 0) for sel in correct_set)
            if total_correct_weight > 0:
                scale = base_score / total_correct_weight
            else:
                scale = 0

            score_earned = 0.0
            for sel in user_set:
                if sel in correct_set:
                    # 选对正确选项，得该选项的权重分值
                    score_earned += option_weights.get(sel, 0) * scale
                else:
                    # 选了错误选项，倒扣该选项权重对应的分值
                    # 错误选项如果在option_weights中有定义则用其权重，否则按平均正确选项分值扣减
                    wrong_deduction = option_weights.get(sel, 0) * scale if sel in option_weights else (base_score / len(correct_set) if correct_set else 0)
                    score_earned -= wrong_deduction
        else:
            # 等权方案：每个正确选项权重 = base_score / 正确选项数
            option_weight = base_score / len(correct_set) if correct_set else 0

            score_earned = 0.0
            for sel in user_set:
                if sel in correct_set:
                    # 选对了正确选项，得对应分值
                    score_earned += option_weight
                else:
                    # 选了错误选项，倒扣对应选项的分值
                    score_earned -= option_weight

        # 保底规则：最低得分为0分
        score_earned = max(0, score_earned)
        is_correct = (user_set == correct_set)
        return (is_correct, round(score_earned, 2))

    return (False, 0)


# ============================================================
# Public routes
# ============================================================



@app.get("/", response_class=HTMLResponse)
async def index(request: Request):
    user = get_current_user(request)
    if user:
        if user["role"] == "admin":
            return RedirectResponse("/admin", status_code=302)
        return RedirectResponse("/student", status_code=302)
    return RedirectResponse("/login", status_code=302)

@app.get("/api/site/login-config")
async def get_login_config():
    """Public endpoint: returns which login methods are available"""
    ss = get_site_settings()
    return JSONResponse({
        "allow_register": ss["allow_register"],
        "allow_account_login": ss["allow_account_login"],
        "allow_student_id_login": ss["allow_student_id_login"]
    })

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    user = get_current_user(request)
    if user:
        return RedirectResponse("/student" if user["role"] == "student" else "/admin", status_code=302)
    return templates.TemplateResponse("login.html", {"request": request})

@app.post("/api/login")
async def api_login(request: Request, response: Response):
    data = await request.json()
    username = data.get("username", "")
    password = data.get("password", "")
    login_type = data.get("login_type", "account")

    # First, verify user credentials
    with get_db() as db:
        if login_type == "student_id":
            user = db.execute("SELECT * FROM users WHERE student_id=?", (username,)).fetchone()
            if not user or not verify_password(password, user["password"]):
                return JSONResponse({"error": "学号或密码错误"}, status_code=401)
        else:
            user = db.execute("SELECT * FROM users WHERE username=?", (username,)).fetchone()
            if not user or not verify_password(password, user["password"]):
                return JSONResponse({"error": "用户名或密码错误"}, status_code=401)

    # Check if this login method is allowed (admin is always allowed to bypass restrictions)
    if user["role"] != "admin":
        ss = get_site_settings()
        if login_type == "student_id" and not ss["allow_student_id_login"]:
            return JSONResponse({"error": "学号登录已被管理员关闭"}, status_code=403)
        if login_type != "student_id" and not ss["allow_account_login"]:
            return JSONResponse({"error": "账号登录已被管理员关闭"}, status_code=403)

    token = make_token(user["id"])
    resp = JSONResponse({"success": True, "role": user["role"]})
    resp.set_cookie("session_token", token, httponly=True, max_age=86400 * 7, samesite="lax")
    return resp

@app.post("/api/register")
async def api_register(request: Request):
    # Check if registration is allowed
    ss = get_site_settings()
    if not ss["allow_register"]:
        return JSONResponse({"error": "管理员已关闭自助注册功能，请联系管理员添加账号"}, status_code=403)
    data = await request.json()
    username = data.get("username", "").strip()
    password = data.get("password", "").strip()
    real_name = data.get("real_name", "").strip()

    if not username or not password:
        return JSONResponse({"error": "用户名和密码不能为空"}, status_code=400)
    if len(password) < 6:
        return JSONResponse({"error": "密码至少6位"}, status_code=400)

    with get_db() as db:
        existing = db.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
        if existing:
            return JSONResponse({"error": "用户名已存在"}, status_code=400)
        db.execute(
            "INSERT INTO users (username, password, real_name, student_id, login_type, role, plain_password) VALUES (?, ?, ?, '', 'account', 'student', ?)",
            (username, hash_password(password), real_name or username, password)
        )

    return JSONResponse({"success": True})

@app.get("/logout")
async def logout(response: Response):
    response = RedirectResponse("/login", status_code=302)
    response.delete_cookie("session_token")
    return response

# ============================================================
# Student routes
# ============================================================
@app.get("/student", response_class=HTMLResponse)
async def student_dashboard(request: Request):
    user = require_login(request)
    if user["role"] == "admin":
        return RedirectResponse("/admin", status_code=302)

    with get_db() as db:
        # Get available exams
        exams_raw = db.execute("""
            SELECT e.*, COUNT(eq.id) as question_count
            FROM exams e
            LEFT JOIN exam_questions eq ON e.id = eq.exam_id
            WHERE e.status = 'active'
            GROUP BY e.id
            ORDER BY e.created_at DESC
        """).fetchall()

        # Recalculate question_count including sub-questions
        exams = []
        for exam in exams_raw:
            exam_dict = dict(exam)
            total_q = 0
            exam_qs = db.execute("""
                SELECT q.type, q.sub_questions
                FROM exam_questions eq
                JOIN questions q ON eq.question_id = q.id
                WHERE eq.exam_id = ?
            """, (exam_dict["id"],)).fetchall()
            for eq in exam_qs:
                if eq["type"] in ("shared", "case"):
                    subs = json.loads(eq["sub_questions"]) if eq["sub_questions"] else []
                    total_q += len(subs) if subs else 1
                else:
                    total_q += 1
            exam_dict["question_count"] = total_q
            exams.append(exam_dict)

        # Get user's exam history
        history = db.execute("""
            SELECT es.*, e.title as exam_title, e.duration, e.total_score, e.pass_score
            FROM exam_sessions es
            JOIN exams e ON es.exam_id = e.id
            WHERE es.user_id = ?
            ORDER BY es.start_time DESC
            LIMIT 20
        """, (user["id"],)).fetchall()

        # Get latest blog posts (published only)
        blog_posts = db.execute("""
            SELECT * FROM blog_posts
            WHERE is_published = 1
            ORDER BY is_pinned DESC, created_at DESC
            LIMIT 10
        """).fetchall()

    return templates.TemplateResponse("dashboard.html", {
        "request": request,
        "user": user,
        "exams": [dict(e) for e in exams],
        "history": [dict(h) for h in history],
        "blog_posts": [dict(p) for p in blog_posts],
    })

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard(request: Request):
    return RedirectResponse("/student", status_code=302)

@app.get("/exam/{exam_id}", response_class=HTMLResponse)
async def exam_page(request: Request, exam_id: int):
    user = require_login(request)
    if user["role"] == "admin":
        return RedirectResponse("/admin", status_code=302)

    with get_db() as db:
        exam = db.execute("SELECT * FROM exams WHERE id=?", (exam_id,)).fetchone()
        if not exam:
            raise HTTPException(status_code=404, detail="考试不存在")

        # Check for existing session (any status)
        existing = db.execute(
            "SELECT * FROM exam_sessions WHERE user_id=? AND exam_id=? ORDER BY id DESC LIMIT 1",
            (user["id"], exam_id)
        ).fetchone()

        review_mode = False
        if existing and existing["status"] == "submitted":
            # Already submitted - enter review mode
            review_mode = True
            session_id = existing["id"]
        elif existing and existing["status"] == "in_progress":
            session_id = existing["id"]
        else:
            # Create new session
            cursor = db.execute(
                "INSERT INTO exam_sessions (user_id, exam_id, status) VALUES (?, ?, 'in_progress')",
                (user["id"], exam_id)
            )
            session_id = cursor.lastrowid

        # Get exam questions
        questions = db.execute("""
            SELECT eq.display_order, eq.score_override, q.*
            FROM exam_questions eq
            JOIN questions q ON eq.question_id = q.id
            WHERE eq.exam_id = ?
            ORDER BY eq.display_order
        """, (exam_id,)).fetchall()

        # Get existing answers
        answers = db.execute("""
            SELECT question_id, sub_index, answer FROM session_answers
            WHERE session_id = ?
        """, (session_id,)).fetchall()

        # Get last position for restoring
        session_row = db.execute(
            "SELECT last_question_index, last_sub_index FROM exam_sessions WHERE id=?",
            (session_id,)
        ).fetchone()
        last_qi = session_row["last_question_index"] if session_row and session_row["last_question_index"] else 0
        last_si = session_row["last_sub_index"] if session_row and session_row["last_sub_index"] else 0

    # Prepare question data for JS
    question_data = []
    answer_map = {}
    for a in answers:
        key = f"{a['question_id']}_{a['sub_index']}"
        answer_map[key] = a["answer"]

    for idx, q in enumerate(questions):
        qd = {
            "id": q["id"],
            "type": q["type"],
            "content": q["content"],
            "options": json.loads(q["options"]) if q["options"] else [],
            "score": q["score_override"] if q["score_override"] is not None else q["score"],
            "display_order": idx + 1,
        }
        if q["type"] == "shared" or q["type"] == "case":
            qd["sub_questions"] = json.loads(q["sub_questions"]) if q["sub_questions"] else []
        # Pre-fill answers
        if q["type"] in ("shared", "case"):
            subs = json.loads(q["sub_questions"]) if q["sub_questions"] else []
            for si in range(len(subs)):
                key = f"{q['id']}_{si}"
                qd.setdefault("user_answers", {})[str(si)] = answer_map.get(key, "")
        else:
            key = f"{q['id']}_0"
            qd["user_answer"] = answer_map.get(key, "")

        # In review mode, attach correct answer and explanation
        if review_mode:
            qd["correct_answer"] = q["answer"]
            qd["explanation"] = q["explanation"]
            if q["type"] in ("shared", "case"):
                subs = json.loads(q["sub_questions"]) if q["sub_questions"] else []
                for si, sub in enumerate(subs):
                    sub["correct_answer"] = sub.get("answer", "")

        question_data.append(qd)

    exam_dict = dict(exam)
    return templates.TemplateResponse("exam.html", {
        "request": request,
        "user": user,
        "exam": exam_dict,
        "session_id": session_id,
        "questions": question_data,
        "question_count": len(question_data),
        "review_mode": review_mode,
        "last_question_index": last_qi,
        "last_sub_index": last_si,
    })

@app.get("/result/{session_id}", response_class=HTMLResponse)
async def result_page(request: Request, session_id: int):
    user = require_login(request)

    with get_db() as db:
        session = db.execute("""
            SELECT es.*, e.title, e.duration, e.total_score, e.pass_score, e.is_practice
            FROM exam_sessions es
            JOIN exams e ON es.exam_id = e.id
            WHERE es.id = ?
        """, (session_id,)).fetchone()

        if not session:
            raise HTTPException(status_code=404, detail="记录不存在")
        if session["user_id"] != user["id"] and user["role"] != "admin":
            raise HTTPException(status_code=403, detail="无权查看")

        # Get answers with question details
        answers = db.execute("""
            SELECT sa.*, q.type, q.content, q.options, q.answer as correct_answer,
                   q.explanation, q.score as base_score, q.sub_questions,
                   eq.display_order, eq.score_override
            FROM session_answers sa
            JOIN questions q ON sa.question_id = q.id
            JOIN exam_questions eq ON eq.exam_id = ? AND eq.question_id = q.id
            WHERE sa.session_id = ?
            ORDER BY eq.display_order, sa.sub_index
        """, (session["exam_id"], session_id)).fetchall()

    # Group answers by question
    question_results = {}
    for a in answers:
        qid = a["question_id"]
        if qid not in question_results:
            question_results[qid] = {
                "id": qid,
                "type": a["type"],
                "content": a["content"],
                "options": json.loads(a["options"]) if a["options"] else [],
                "correct_answer": a["correct_answer"],
                "explanation": a["explanation"],
                "display_order": a["display_order"],
                "score": a["score_override"] if a["score_override"] is not None else a["base_score"],
                "earned_score": 0,
                "is_correct": True,
                "user_answers": {},
            }
            if a["type"] in ("shared", "case"):
                question_results[qid]["sub_questions"] = json.loads(a["sub_questions"]) if a["sub_questions"] else []

        si = a["sub_index"]
        question_results[qid]["user_answers"][str(si)] = a["answer"]
        if not a["is_correct"]:
            question_results[qid]["is_correct"] = False
        question_results[qid]["earned_score"] += a["score"]

    # Normalize case-type user answers: JSON array -> comma-separated string
    for qr in question_results.values():
        if qr["type"] == "case":
            for si, ans in list(qr["user_answers"].items()):
                if ans:
                    try:
                        parsed = json.loads(ans)
                        if isinstance(parsed, list):
                            qr["user_answers"][si] = ",".join(parsed)
                    except (json.JSONDecodeError, TypeError):
                        pass

    return templates.TemplateResponse("result.html", {
        "request": request,
        "user": user,
        "session": dict(session),
        "results": list(question_results.values()),
    })

# ============================================================
# Student API
# ============================================================
@app.post("/api/exam/{session_id}/save")
async def save_answer(session_id: int, request: Request):
    user = require_login(request)
    data = await request.json()
    question_id = data.get("question_id")
    sub_index = data.get("sub_index", 0)
    answer = data.get("answer", "")
    current_question_index = data.get("current_question_index")
    current_sub_index = data.get("current_sub_index")

    with get_db() as db:
        # Verify session belongs to user
        session = db.execute(
            "SELECT id FROM exam_sessions WHERE id=? AND user_id=? AND status='in_progress'",
            (session_id, user["id"])
        ).fetchone()
        if not session:
            return JSONResponse({"error": "考试会话不存在"}, status_code=404)

        # Upsert answer
        existing = db.execute(
            "SELECT id FROM session_answers WHERE session_id=? AND question_id=? AND sub_index=?",
            (session_id, question_id, sub_index)
        ).fetchone()

        if existing:
            db.execute(
                "UPDATE session_answers SET answer=? WHERE id=?",
                (answer, existing["id"])
            )
        else:
            db.execute(
                "INSERT INTO session_answers (session_id, question_id, sub_index, answer) VALUES (?, ?, ?, ?)",
                (session_id, question_id, sub_index, answer)
            )

        # Update last position
        if current_question_index is not None:
            db.execute(
                "UPDATE exam_sessions SET last_question_index=?, last_sub_index=? WHERE id=?",
                (current_question_index, current_sub_index or 0, session_id)
            )

    return JSONResponse({"success": True})

@app.post("/api/exam/{session_id}/submit")
async def submit_exam(session_id: int, request: Request):
    user = require_login(request)

    with get_db() as db:
        session = db.execute(
            "SELECT * FROM exam_sessions WHERE id=? AND user_id=?",
            (session_id, user["id"])
        ).fetchone()
        if not session:
            return JSONResponse({"error": "考试会话不存在"}, status_code=404)
        if session["status"] == "submitted":
            return JSONResponse({"error": "考试已提交"}, status_code=400)

        exam = db.execute("SELECT * FROM exams WHERE id=?", (session["exam_id"],)).fetchone()

        # Get all questions for this exam
        exam_qs = db.execute("""
            SELECT eq.question_id, eq.score_override, q.type, q.answer, q.score, q.sub_questions
            FROM exam_questions eq
            JOIN questions q ON eq.question_id = q.id
            WHERE eq.exam_id = ?
        """, (session["exam_id"],)).fetchall()

        # Get all answers
        answers = db.execute("""
            SELECT question_id, sub_index, answer FROM session_answers WHERE session_id=?
        """, (session_id,)).fetchall()

        answer_map = {}
        for a in answers:
            answer_map[(a["question_id"], a["sub_index"])] = a["answer"]

        # Score each question
        total_score = 0.0
        for eq in exam_qs:
            qid = eq["question_id"]
            q_type = eq["type"]
            correct_answer = eq["answer"]
            base_score = eq["score_override"] if eq["score_override"] is not None else eq["score"]

            if q_type in ("shared", "case"):
                subs = json.loads(eq["sub_questions"]) if eq["sub_questions"] else []
                for si, sub in enumerate(subs):
                    user_ans = answer_map.get((qid, si), "")
                    sub_correct = sub.get("answer", "")
                    sub_score = 1.0  # 每个提问1个得分点

                    if q_type == "case":
                        # Extract option weights from sub-question options
                        sub_opts = sub.get("options", [])
                        opt_weights = {}
                        for opt in sub_opts:
                            if isinstance(opt, dict) and "weight" in opt and opt.get("weight") is not None:
                                opt_weights[opt.get("label", "").upper()] = float(opt["weight"])
                        is_correct, earned = score_question("case", user_ans, sub_correct, sub_score, opt_weights if opt_weights else None)
                    elif q_type == "shared":
                        is_correct, earned = score_question("single", user_ans, sub_correct, sub_score)
                    else:
                        is_correct, earned = False, 0

                    total_score += earned

                    # Update answer record
                    existing = db.execute(
                        "SELECT id FROM session_answers WHERE session_id=? AND question_id=? AND sub_index=?",
                        (session_id, qid, si)
                    ).fetchone()
                    if existing:
                        db.execute(
                            "UPDATE session_answers SET is_correct=?, score=? WHERE id=?",
                            (is_correct, earned, existing["id"])
                        )
                    else:
                        db.execute(
                            "INSERT INTO session_answers (session_id, question_id, sub_index, answer, is_correct, score) VALUES (?,?,?,?,?,?)",
                            (session_id, qid, si, user_ans, is_correct, earned)
                        )
            else:
                user_ans = answer_map.get((qid, 0), "")
                is_correct, earned = score_question(q_type, user_ans, correct_answer, base_score)
                total_score += earned

                existing = db.execute(
                    "SELECT id FROM session_answers WHERE session_id=? AND question_id=? AND sub_index=0",
                    (session_id, qid)
                ).fetchone()
                if existing:
                    db.execute(
                        "UPDATE session_answers SET is_correct=?, score=? WHERE id=?",
                        (is_correct, earned, existing["id"])
                    )
                else:
                    db.execute(
                        "INSERT INTO session_answers (session_id, question_id, sub_index, answer, is_correct, score) VALUES (?,?,?,?,?,?)",
                        (session_id, qid, 0, user_ans, is_correct, earned)
                    )

        total_score = round(total_score, 2)
        now = datetime.now().isoformat()
        db.execute(
            "UPDATE exam_sessions SET status='submitted', end_time=?, total_score=? WHERE id=?",
            (now, total_score, session_id)
        )

    return JSONResponse({"success": True, "score": total_score, "session_id": session_id})

# ============================================================
# Admin routes
# ============================================================
@app.get("/admin", response_class=HTMLResponse)
async def admin_page(request: Request):
    user = require_admin(request)
    with get_db() as db:
        stats = {
            "user_count": db.execute("SELECT COUNT(*) as c FROM users WHERE role='student'").fetchone()["c"],
            "question_count": db.execute("SELECT COUNT(*) as c FROM questions").fetchone()["c"],
            "exam_count": db.execute("SELECT COUNT(*) as c FROM exams").fetchone()["c"],
            "session_count": db.execute("SELECT COUNT(*) as c FROM exam_sessions WHERE status='submitted'").fetchone()["c"],
            "avg_score": db.execute("SELECT COALESCE(AVG(total_score), 0) as avg FROM exam_sessions WHERE status='submitted'").fetchone()["avg"],
        }
        recent_sessions = db.execute("""
            SELECT es.*, u.real_name, u.username, e.title as exam_title
            FROM exam_sessions es
            JOIN users u ON es.user_id = u.id
            JOIN exams e ON es.exam_id = e.id
            WHERE es.status = 'submitted'
            ORDER BY es.end_time DESC LIMIT 10
        """).fetchall()

    return templates.TemplateResponse("admin.html", {
        "request": request,
        "user": user,
        "stats": stats,
        "recent_sessions": [dict(s) for s in recent_sessions],
    })

# ============================================================
# Admin API - Questions
# ============================================================
@app.get("/api/admin/questions")
async def list_questions(request: Request, q_type: str = None, category: str = None, source_set: str = None, page: int = 1, size: int = 20):
    user = require_admin(request)
    with get_db() as db:
        conditions = []
        params = []
        if q_type:
            conditions.append("type=?")
            params.append(q_type)
        if category:
            if category in ("总论", "各论"):
                conditions.append("category LIKE ?")
                params.append(f"{category}%")
            else:
                conditions.append("category=?")
                params.append(category)
        if source_set:
            conditions.append("source_set=?")
            params.append(source_set)
        where = " WHERE " + " AND ".join(conditions) if conditions else ""
        query = f"SELECT * FROM questions{where} ORDER BY created_at DESC"
        query += f" LIMIT {size} OFFSET {(page-1)*size}"

        questions = db.execute(query, params).fetchall()
        total = db.execute(f"SELECT COUNT(*) as c FROM questions{where}", params).fetchone()["c"]

    return JSONResponse({
        "items": [dict(q) for q in questions],
        "total": total,
        "page": page,
        "size": size,
    })

@app.post("/api/admin/questions")
async def create_question(request: Request):
    user = require_admin(request)
    data = await request.json()

    q_type = data.get("type")
    content = data.get("content", "")
    options = data.get("options", [])
    answer = data.get("answer", "")
    explanation = data.get("explanation", "")
    category = data.get("category", "")
    score = data.get("score", 1)
    sub_questions = data.get("sub_questions", [])
    images = data.get("images", [])

    if not content:
        return JSONResponse({"error": "题目内容不能为空"}, status_code=400)
    if q_type in ("single", "multiple") and not answer:
        return JSONResponse({"error": "答案不能为空"}, status_code=400)

    with get_db() as db:
        db.execute(
            """INSERT INTO questions (type, category, content, options, answer, explanation, score, sub_questions, images)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (q_type, category, content, json.dumps(options, ensure_ascii=False),
             answer, explanation, score, json.dumps(sub_questions, ensure_ascii=False),
             json.dumps(images, ensure_ascii=False))
        )

    return JSONResponse({"success": True})

@app.put("/api/admin/questions/{qid}")
async def update_question(qid: int, request: Request):
    user = require_admin(request)
    data = await request.json()

    with get_db() as db:
        db.execute("""
            UPDATE questions SET type=?, category=?, content=?, options=?, answer=?,
            explanation=?, score=?, sub_questions=?, images=? WHERE id=?
        """, (
            data.get("type"), data.get("category", ""), data.get("content", ""),
            json.dumps(data.get("options", []), ensure_ascii=False),
            data.get("answer", ""), data.get("explanation", ""),
            data.get("score", 1),
            json.dumps(data.get("sub_questions", []), ensure_ascii=False),
            json.dumps(data.get("images", []), ensure_ascii=False),
            qid
        ))

    return JSONResponse({"success": True})

@app.delete("/api/admin/questions/{qid}")
async def delete_question(qid: int, request: Request):
    user = require_admin(request)
    with get_db() as db:
        db.execute("DELETE FROM questions WHERE id=?", (qid,))
    return JSONResponse({"success": True})

@app.post("/api/admin/questions/batch-delete")
async def batch_delete_questions(request: Request):
    user = require_admin(request)
    body = await request.json()
    ids = body.get("ids", [])
    if not ids:
        return JSONResponse({"error": "未选择任何题目"}, status_code=400)
    with get_db() as db:
        placeholders = ",".join("?" * len(ids))
        db.execute(f"DELETE FROM questions WHERE id IN ({placeholders})", ids)
    return JSONResponse({"success": True, "deleted": len(ids)})

# ============================================================
# Admin API - Import Questions
# ============================================================
@app.post("/api/admin/import/json")
async def import_json(request: Request, file: UploadFile = File(...), source_set: str = Form("")):
    user = require_admin(request)
    content = await file.read()

    try:
        data = json.loads(content.decode("utf-8"))
    except (json.JSONDecodeError, UnicodeDecodeError):
        return JSONResponse({"error": "JSON格式错误"}, status_code=400)

    if not isinstance(data, list):
        return JSONResponse({"error": "JSON必须为数组"}, status_code=400)

    count = 0
    errors = []
    question_ids = []
    with get_db() as db:
        for idx, item in enumerate(data):
            try:
                q_type = item.get("type", "")
                if q_type not in ("single", "multiple", "shared", "case"):
                    errors.append(f"第{idx+1}题: 无效题型 {q_type}")
                    continue

                cursor = db.execute(
                    """INSERT INTO questions (type, category, content, options, answer, explanation, score, sub_questions, source_set)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (
                        q_type,
                        item.get("category", ""),
                        item.get("content", ""),
                        json.dumps(item.get("options", []), ensure_ascii=False),
                        item.get("answer", ""),
                        item.get("explanation", ""),
                        item.get("score", 1),
                        json.dumps(item.get("sub_questions", []), ensure_ascii=False),
                        source_set or item.get("source_set", ""),
                    )
                )
                question_ids.append(cursor.lastrowid)
                count += 1
            except Exception as e:
                errors.append(f"第{idx+1}题: {str(e)}")

        # Auto-create exam if source_set is provided
        exam_id = auto_create_exam_from_source_set(db, source_set, question_ids)

    result = {"success": True, "imported": count, "errors": errors, "source_set": source_set}
    if exam_id:
        result["exam_id"] = exam_id
        result["message"] = f"已自动创建考试「{source_set}」"
    return JSONResponse(result)

@app.post("/api/admin/import/excel")
async def import_excel(request: Request, file: UploadFile = File(...), source_set: str = Form("")):
    user = require_admin(request)
    content = await file.read()

    try:
        import openpyxl
        from io import BytesIO
        wb = openpyxl.load_workbook(BytesIO(content))
        ws = wb.active

        count = 0
        errors = []
        question_ids = []
        with get_db() as db:
            for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
                try:
                    if not row or not row[0]:
                        continue
                    q_type = str(row[0]).strip()
                    category = str(row[1]).strip() if row[1] else ""
                    q_content = str(row[2]).strip() if row[2] else ""
                    options_str = str(row[3]).strip() if row[3] else "[]"
                    answer = str(row[4]).strip() if row[4] else ""
                    explanation = str(row[5]).strip() if row[5] else ""
                    score_val = float(row[6]) if row[6] else 1
                    sub_q_str = str(row[7]).strip() if len(row) > 7 and row[7] else "[]"

                    # Parse options
                    try:
                        options = json.loads(options_str) if options_str.startswith("[") else []
                        if not options:
                            import re
                            parts = re.split(r'[A-H][\.、．]\s*', options_str)
                            labels = re.findall(r'([A-H])[\.、．]', options_str)
                            options = [{"label": l, "text": p.strip()} for l, p in zip(labels, parts) if p.strip()]
                    except json.JSONDecodeError:
                        options = []

                    # Parse sub_questions
                    try:
                        sub_questions = json.loads(sub_q_str) if sub_q_str.startswith("[") else []
                    except json.JSONDecodeError:
                        sub_questions = []

                    if q_type not in ("single", "multiple", "shared", "case"):
                        errors.append(f"第{row_idx}行: 无效题型")
                        continue

                    cursor = db.execute(
                        """INSERT INTO questions (type, category, content, options, answer, explanation, score, sub_questions, source_set)
                           VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                        (q_type, category, q_content, json.dumps(options, ensure_ascii=False),
                         answer, explanation, score_val, json.dumps(sub_questions, ensure_ascii=False),
                         source_set)
                    )
                    question_ids.append(cursor.lastrowid)
                    count += 1
                except Exception as e:
                    errors.append(f"第{row_idx}行: {str(e)}")

            # Auto-create exam if source_set is provided
            exam_id = auto_create_exam_from_source_set(db, source_set, question_ids)

        result = {"success": True, "imported": count, "errors": errors, "source_set": source_set}
        if exam_id:
            result["exam_id"] = exam_id
            result["message"] = f"已自动创建考试「{source_set}」"
        return JSONResponse(result)
    except Exception as e:
        return JSONResponse({"error": f"Excel解析失败: {str(e)}"}, status_code=400)

@app.get("/api/admin/template/excel")
async def download_template(request: Request):
    user = require_admin(request)
    import openpyxl
    from io import BytesIO

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "试题导入模板"

    headers = ["题型", "分类", "题目内容", "选项(JSON)", "答案", "解析", "分值", "子题(JSON)"]
    ws.append(headers)

    # Example rows
    ws.append([
        "single", "总论",
        "细胞坏死的主要形态学标志是",
        json.dumps([{"label":"A","text":"细胞膜破裂"},{"label":"B","text":"线粒体肿胀"},{"label":"C","text":"内质网扩张"},{"label":"D","text":"核固缩"},{"label":"E","text":"高尔基体解体"}], ensure_ascii=False),
        "A", "细胞坏死时细胞膜破裂是最重要的形态学标志", 1, "[]"
    ])
    ws.append([
        "multiple", "总论",
        "下列哪些属于适应性改变",
        json.dumps([{"label":"A","text":"萎缩"},{"label":"B","text":"肥大"},{"label":"C","text":"增生"},{"label":"D","text":"化生"},{"label":"E","text":"坏死"}], ensure_ascii=False),
        "ABCD", "适应性改变包括萎缩、肥大、增生和化生", 1, "[]"
    ])
    ws.append([
        "shared", "各论",
        "患者男性，65岁，吸烟史40年，慢性咳嗽咳痰20年，近5年出现活动后气短。",
        "[]", "", "共用题干题", 1,
        json.dumps([
            {"content": "最可能的诊断是", "answer": "C", "options": [{"label":"A","text":"支气管哮喘"},{"label":"B","text":"支气管扩张"},{"label":"C","text":"慢性阻塞性肺疾病"},{"label":"D","text":"肺结核"},{"label":"E","text":"肺癌"}]},
            {"content": "确诊首选的检查是", "answer": "B", "options": [{"label":"A","text":"胸部CT"},{"label":"B","text":"肺功能检查"},{"label":"C","text":"痰培养"},{"label":"D","text":"血气分析"},{"label":"E","text":"支气管镜"}]}
        ], ensure_ascii=False)
    ])
    ws.append([
        "case", "各论",
        "患者女性，45岁，发现右乳肿块2个月。查体：右乳外上象限可及一3cm×2cm肿块，质硬，边界不清，活动度差。同侧腋窝可触及肿大淋巴结。",
        "[]", "", "案例分析题", 3,
        json.dumps([
            {"content": "最可能的诊断", "options": [{"label":"A","text":"乳腺纤维腺瘤"},{"label":"B","text":"乳腺癌"},{"label":"C","text":"乳腺囊肿"},{"label":"D","text":"乳腺增生"},{"label":"E","text":"乳腺炎"},{"label":"F","text":"脂肪瘤"}], "answer": "B"},
            {"content": "为明确诊断应首选的检查", "options": [{"label":"A","text":"B超"},{"label":"B","text":"钼靶X线"},{"label":"C","text":"穿刺活检"},{"label":"D","text":"MRI"},{"label":"E","text":"PET-CT"},{"label":"F","text":"CT"}], "answer": "C"},
            {"content": "若确诊为乳腺癌，最常见的病理类型是", "options": [{"label":"A","text":"浸润性导管癌"},{"label":"B","text":"浸润性小叶癌"},{"label":"C","text":"髓样癌"},{"label":"D","text":"黏液癌"},{"label":"E","text":"Paget病"},{"label":"F","text":"小管癌"}], "answer": "A"}
        ], ensure_ascii=False)
    ])

    # Set column widths
    widths = [10, 10, 40, 60, 10, 40, 8, 60]
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=question_template.xlsx"}
    )

# ============================================================
# Admin API - Exams
# ============================================================
@app.get("/api/admin/exams")
async def list_exams(request: Request):
    user = require_admin(request)
    with get_db() as db:
        exams = db.execute("""
            SELECT e.*, COUNT(eq.id) as question_count
            FROM exams e
            LEFT JOIN exam_questions eq ON e.id = eq.exam_id
            GROUP BY e.id
            ORDER BY e.created_at DESC
        """).fetchall()
    return JSONResponse({"items": [dict(e) for e in exams]})

@app.get("/api/admin/exams/{eid}/questions")
async def get_exam_questions(eid: int, request: Request):
    user = require_admin(request)
    with get_db() as db:
        rows = db.execute(
            """SELECT eq.question_id, eq.score_override, q.type
                FROM exam_questions eq
                JOIN questions q ON eq.question_id = q.id
                WHERE eq.exam_id=? ORDER BY eq.display_order""",
            (eid,)
        ).fetchall()
    question_ids = [r["question_id"] for r in rows]
    # Build type score map from score_overrides
    type_scores = {}
    for r in rows:
        if r["score_override"] is not None and r["type"]:
            type_scores[r["type"]] = r["score_override"]
    return JSONResponse({"question_ids": question_ids, "type_scores": type_scores})

@app.post("/api/admin/exams")
async def create_exam(request: Request):
    user = require_admin(request)
    data = await request.json()

    with get_db() as db:
        cursor = db.execute(
            """INSERT INTO exams (title, description, duration, total_score, pass_score, is_practice, shuffle_questions, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (
                data.get("title", ""),
                data.get("description", ""),
                data.get("duration", 120),
                data.get("total_score", 100),
                data.get("pass_score", 60),
                data.get("is_practice", False),
                data.get("shuffle_questions", False),
                data.get("status", "active"),
            )
        )
        exam_id = cursor.lastrowid

        # Add questions
        question_ids = data.get("question_ids", [])
        type_scores = data.get("type_scores", {})
        for idx, qid in enumerate(question_ids):
            score_ov = None
            if type_scores:
                q_info = db.execute("SELECT type FROM questions WHERE id=?", (qid,)).fetchone()
                if q_info and q_info["type"] in type_scores:
                    score_ov = type_scores[q_info["type"]]
            db.execute(
                "INSERT INTO exam_questions (exam_id, question_id, display_order, score_override) VALUES (?, ?, ?, ?)",
                (exam_id, qid, idx + 1, score_ov)
            )

    return JSONResponse({"success": True, "exam_id": exam_id})

@app.put("/api/admin/exams/{eid}")
async def update_exam(eid: int, request: Request):
    user = require_admin(request)
    data = await request.json()

    with get_db() as db:
        db.execute("""
            UPDATE exams SET title=?, description=?, duration=?, total_score=?, pass_score=?,
            is_practice=?, shuffle_questions=?, status=? WHERE id=?
        """, (
            data.get("title"), data.get("description", ""),
            data.get("duration", 120), data.get("total_score", 100),
            data.get("pass_score", 60), data.get("is_practice", False),
            data.get("shuffle_questions", False), data.get("status", "active"),
            eid
        ))

        # Update questions if provided
        if "question_ids" in data:
            db.execute("DELETE FROM exam_questions WHERE exam_id=?", (eid,))
            type_scores = data.get("type_scores", {})
            for idx, qid in enumerate(data["question_ids"]):
                score_ov = None
                if type_scores:
                    q_info = db.execute("SELECT type FROM questions WHERE id=?", (qid,)).fetchone()
                    if q_info and q_info["type"] in type_scores:
                        score_ov = type_scores[q_info["type"]]
                db.execute(
                    "INSERT INTO exam_questions (exam_id, question_id, display_order, score_override) VALUES (?, ?, ?, ?)",
                    (eid, qid, idx + 1, score_ov)
                )

    return JSONResponse({"success": True})

@app.delete("/api/admin/exams/{eid}")
async def delete_exam(eid: int, request: Request):
    user = require_admin(request)
    with get_db() as db:
        # Cascade delete: session_answers -> exam_sessions -> exam_questions -> exam
        db.execute("""DELETE FROM session_answers WHERE session_id IN
                      (SELECT id FROM exam_sessions WHERE exam_id=?)""", (eid,))
        db.execute("DELETE FROM exam_sessions WHERE exam_id=?", (eid,))
        db.execute("DELETE FROM exam_questions WHERE exam_id=?", (eid,))
        db.execute("DELETE FROM exams WHERE id=?", (eid,))
    return JSONResponse({"success": True})

@app.post("/api/admin/exams/batch-delete")
async def batch_delete_exams(request: Request):
    user = require_admin(request)
    body = await request.json()
    ids = body.get("ids", [])
    if not ids:
        return JSONResponse({"error": "未选择任何考试"}, status_code=400)
    with get_db() as db:
        placeholders = ",".join("?" * len(ids))
        db.execute(f"DELETE FROM session_answers WHERE session_id IN (SELECT id FROM exam_sessions WHERE exam_id IN ({placeholders}))", ids)
        db.execute(f"DELETE FROM exam_sessions WHERE exam_id IN ({placeholders})", ids)
        db.execute(f"DELETE FROM exam_questions WHERE exam_id IN ({placeholders})", ids)
        db.execute(f"DELETE FROM exams WHERE id IN ({placeholders})", ids)
    return JSONResponse({"success": True, "deleted": len(ids)})



@app.get("/api/admin/exams/{eid}/export")
async def export_exam(eid: int, request: Request):
    """Export exam as DOCX file with questions and answers."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import io, urllib.parse

    user = require_admin(request)
    with get_db() as db:
        exam = db.execute("SELECT * FROM exams WHERE id=?", (eid,)).fetchone()
        if not exam:
            return JSONResponse({"error": "试卷不存在"}, status_code=404)

        rows = db.execute(
            """SELECT q.id, q.type, q.content, q.options, q.answer, q.explanation, q.sub_questions,
                      eq.display_order, eq.score_override
               FROM exam_questions eq
               JOIN questions q ON eq.question_id = q.id
               WHERE eq.exam_id=? ORDER BY eq.display_order""",
            (eid,)
        ).fetchall()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')

    # Title
    title_para = doc.add_heading(exam["title"], level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title_para.runs:
        run.font.size = Pt(22)

    # Info line
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_run = info.add_run(f"考试时长：{exam['duration']}分钟　　满分：{exam['total_score']}分　　合格线：{exam['pass_score']}分")
    info_run.font.size = Pt(10)
    info_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    type_names = {"single": "单选题", "multiple": "多选题", "shared": "共用题干单选题", "case": "案例分析题"}

    q_num = 0
    for row in rows:
        q_type = row["type"]
        content_text = row["content"] or ""
        options_raw = row["options"]
        answer = row["answer"] or ""
        explanation = row["explanation"] or ""
        sub_questions = row["sub_questions"]

        if q_type in ("single", "multiple"):
            q_num += 1
            type_hint = f"（{type_names.get(q_type, q_type)}）" if q_type == "multiple" else ""

            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{q_num}. {content_text}{type_hint}")
            q_run.bold = True

            if options_raw:
                try:
                    opts = json.loads(options_raw) if isinstance(options_raw, str) else options_raw
                    for opt in opts:
                        opt_para = doc.add_paragraph()
                        opt_para.paragraph_format.left_indent = Cm(1)
                        opt_para.add_run(f"{opt['label']}. {opt['text']}")
                except:
                    pass

            ans_para = doc.add_paragraph()
            ans_run = ans_para.add_run(f"答案：{answer}")
            ans_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

            if explanation:
                exp_para = doc.add_paragraph()
                exp_run = exp_para.add_run(f"解析：{explanation}")
                exp_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                exp_run.font.size = Pt(10)

            doc.add_paragraph()

        elif q_type in ("shared", "case"):
            q_num += 1
            type_hint = type_names.get(q_type, q_type)

            q_para = doc.add_paragraph()
            q_run = q_para.add_run(f"{q_num}. 【{type_hint}】{content_text}")
            q_run.bold = True

            if sub_questions:
                try:
                    subs = json.loads(sub_questions) if isinstance(sub_questions, str) else sub_questions
                    for si, sub in enumerate(subs):
                        sub_content = sub.get("content", "")
                        sub_answer = sub.get("answer", "")
                        sub_options = sub.get("options", [])
                        sub_explanation = sub.get("explanation", "")

                        sub_para = doc.add_paragraph()
                        sub_run = sub_para.add_run(f"  提问{si+1}：{sub_content}")
                        sub_run.bold = True

                        if sub_options:
                            try:
                                opts = json.loads(sub_options) if isinstance(sub_options, str) else sub_options
                                for opt in opts:
                                    opt_para = doc.add_paragraph()
                                    opt_para.paragraph_format.left_indent = Cm(2)
                                    opt_para.add_run(f"{opt['label']}. {opt['text']}")
                            except:
                                pass

                        ans_para = doc.add_paragraph()
                        ans_run = ans_para.add_run(f"  答案：{sub_answer}")
                        ans_run.font.color.rgb = RGBColor(0x00, 0x80, 0x00)

                        if sub_explanation:
                            exp_para = doc.add_paragraph()
                            exp_run = exp_para.add_run(f"  解析：{sub_explanation}")
                            exp_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                            exp_run.font.size = Pt(10)

                        doc.add_paragraph()
                except:
                    pass

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    filename = f"{exam['title']}.docx"
    encoded_filename = urllib.parse.quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


# ============================================================
# Admin API - Users
# ============================================================
@app.get("/api/admin/users")
async def list_users(request: Request):
    user = require_admin(request)
    search = request.query_params.get("search", "").strip()
    with get_db() as db:
        if search:
            users = db.execute("""
                SELECT u.*, COUNT(es.id) as exam_count,
                       COALESCE(AVG(CASE WHEN es.status='submitted' THEN es.total_score END), 0) as avg_score
                FROM users u
                LEFT JOIN exam_sessions es ON u.id = es.user_id
                WHERE u.role = 'student' AND (u.username LIKE ? OR u.student_id LIKE ? OR u.real_name LIKE ?)
                GROUP BY u.id
                ORDER BY u.created_at DESC
            """, ("%" + search + "%", "%" + search + "%", "%" + search + "%")).fetchall()
        else:
            users = db.execute("""
                SELECT u.*, COUNT(es.id) as exam_count,
                       COALESCE(AVG(CASE WHEN es.status='submitted' THEN es.total_score END), 0) as avg_score
                FROM users u
                LEFT JOIN exam_sessions es ON u.id = es.user_id
                WHERE u.role = 'student'
                GROUP BY u.id
                ORDER BY u.created_at DESC
            """).fetchall()
    items = []
    for u in users:
        d = dict(u)
        d.pop("password", None)
        items.append(d)
    return JSONResponse({"items": items, "total": len(items)})

@app.post("/api/admin/users")
async def create_user(request: Request):
    user = require_admin(request)
    data = await request.json()

    username = data.get("username", "").strip()
    password = data.get("password", "").strip()
    real_name = data.get("real_name", "").strip()
    student_id = data.get("student_id", "").strip()
    login_type = "student_id" if student_id else "account"
    role = data.get("role", "student")

    if not username or not password:
        return JSONResponse({"error": "用户名和密码不能为空"}, status_code=400)

    with get_db() as db:
        existing = db.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
        if existing:
            return JSONResponse({"error": "用户名已存在"}, status_code=400)
        if student_id:
            existing_sid = db.execute("SELECT id FROM users WHERE student_id=?", (student_id,)).fetchone()
            if existing_sid:
                return JSONResponse({"error": "学号已存在"}, status_code=400)
        db.execute(
            "INSERT INTO users (username, password, real_name, student_id, login_type, role, plain_password) VALUES (?, ?, ?, ?, ?, ?, ?)",
            (username, hash_password(password), real_name or username, student_id, login_type, role, password)
        )

    return JSONResponse({"success": True})

@app.delete("/api/admin/users/{uid}")
async def delete_user(uid: int, request: Request):
    user = require_admin(request)
    with get_db() as db:
        # Check user exists and protect admin account
        row = db.execute("SELECT role, username FROM users WHERE id=?", (uid,)).fetchone()
        if not row:
            return JSONResponse({"error": "用户不存在"}, status_code=404)
        if row["role"] == "admin" or row["username"] == "admin":
            return JSONResponse({"error": "admin 管理员账号不可删除"}, status_code=400)
        # Delete in correct order: answers -> sessions -> user
        db.execute("DELETE FROM session_answers WHERE session_id IN (SELECT id FROM exam_sessions WHERE user_id=?)", (uid,))
        db.execute("DELETE FROM exam_sessions WHERE user_id=?", (uid,))
        db.execute("DELETE FROM users WHERE id=?", (uid,))
    return JSONResponse({"success": True})

@app.post("/api/admin/users/batch-delete")
async def batch_delete_users(request: Request):
    user = require_admin(request)
    body = await request.json()
    ids = body.get("ids", [])
    if not ids:
        return JSONResponse({"error": "未选择任何用户"}, status_code=400)
    # Protect admin users from being deleted
    with get_db() as db:
        placeholders = ",".join("?" * len(ids))
        admin_ids = [r["id"] for r in db.execute(f"SELECT id FROM users WHERE id IN ({placeholders}) AND role='admin'", ids).fetchall()]
        if admin_ids:
            return JSONResponse({"error": "不能删除管理员账号，请先取消选择管理员"}, status_code=400)
        db.execute(f"DELETE FROM session_answers WHERE session_id IN (SELECT id FROM exam_sessions WHERE user_id IN ({placeholders}))", ids)
        db.execute(f"DELETE FROM exam_sessions WHERE user_id IN ({placeholders})", ids)
        db.execute(f"DELETE FROM users WHERE id IN ({placeholders})", ids)
    return JSONResponse({"success": True, "deleted": len(ids)})

@app.post("/api/admin/users/{uid}/reset-password")
async def reset_password(uid: int, request: Request):
    user = require_admin(request)
    data = await request.json()
    new_password = data.get("new_password", "123456")

    with get_db() as db:
        db.execute("UPDATE users SET password=?, plain_password=? WHERE id=?", (hash_password(new_password), uid))
    return JSONResponse({"success": True})



@app.get("/api/admin/site/settings")
async def get_site_settings_api(request: Request):
    user = require_admin(request)
    return JSONResponse(get_site_settings())

@app.post("/api/admin/site/settings")
async def save_site_settings_api(request: Request):
    user = require_admin(request)
    data = await request.json()
    allow_register = 1 if data.get("allow_register") else 0
    allow_account_login = 1 if data.get("allow_account_login") else 0
    allow_student_id_login = 1 if data.get("allow_student_id_login") else 0
    with get_db() as db:
        db.execute(
            "UPDATE site_settings SET allow_register=?, allow_account_login=?, allow_student_id_login=? WHERE id=1",
            (allow_register, allow_account_login, allow_student_id_login)
        )
    return JSONResponse({"success": True})

@app.get("/api/admin/users/excel-template")
async def user_excel_template(request: Request):
    """下载用户导入Excel模板"""
    user = require_admin(request)
    from openpyxl import Workbook
    from io import BytesIO

    wb = Workbook()
    ws = wb.active
    ws.title = "用户导入模板"
    ws.append(["学号", "姓名", "密码"])
    # 示例数据
    ws.append(["2024001", "张三", "123456"])
    ws.append(["2024002", "李四", "123456"])

    # 设置列宽
    ws.column_dimensions["A"].width = 15
    ws.column_dimensions["B"].width = 15
    ws.column_dimensions["C"].width = 15

    output = BytesIO()
    wb.save(output)
    output.seek(0)

    from fastapi.responses import StreamingResponse
    headers = {"Content-Disposition": "attachment; filename=user_import_template.xlsx"}
    return StreamingResponse(output, media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", headers=headers)


@app.post("/api/admin/users/import-excel")
async def import_users_excel(request: Request, file: UploadFile = File(...)):
    """从Excel批量导入用户"""
    user = require_admin(request)
    content = await file.read()

    try:
        from openpyxl import load_workbook
        from io import BytesIO
        wb = load_workbook(BytesIO(content))
        ws = wb.active
    except Exception as e:
        return JSONResponse({"error": f"文件读取失败: {str(e)}"}, status_code=400)

    # 读取表头，确认列位置
    headers = [str(cell.value or "").strip() for cell in ws[1]]
    col_map = {}
    for i, h in enumerate(headers):
        hl = h.lower()
        if h in ("学号", "student_id", "username", "用户名"):
            col_map["username"] = i
        elif h in ("姓名", "name", "real_name"):
            col_map["real_name"] = i
        elif h in ("密码", "password", "pwd"):
            col_map["password"] = i

    if "username" not in col_map or "password" not in col_map:
        return JSONResponse({"error": "Excel表头必须包含'学号'和'密码'列"}, status_code=400)

    imported = 0
    errors = []
    with get_db() as db:
        for row_idx, row in enumerate(ws.iter_rows(min_row=2, values_only=True), start=2):
            if not row or all(v is None for v in row):
                continue
            cells = [str(v).strip() if v is not None else "" for v in row]

            username = cells[col_map["username"]] if "username" in col_map else ""
            real_name = cells[col_map["real_name"]] if "real_name" in col_map else ""
            password = cells[col_map["password"]] if "password" in col_map else ""

            if not username:
                errors.append(f"第{row_idx}行: 学号为空，已跳过")
                continue
            if not password:
                errors.append(f"第{row_idx}行({username}): 密码为空，已跳过")
                continue

            try:
                existing = db.execute("SELECT id FROM users WHERE username=?", (username,)).fetchone()
                if existing:
                    errors.append(f"第{row_idx}行({username}): 用户名已存在，已跳过")
                    continue
                # Check if student_id already exists
                existing_sid = db.execute("SELECT id FROM users WHERE student_id=?", (username,)).fetchone()
                if existing_sid:
                    errors.append(f"第{row_idx}行({username}): 学号已存在，已跳过")
                    continue
                # Generate unique username from student_id
                gen_username = "stu_" + username
                existing_uname = db.execute("SELECT id FROM users WHERE username=?", (gen_username,)).fetchone()
                if existing_uname:
                    gen_username = "stu_" + username + "_" + str(row_idx)
                db.execute(
                    "INSERT INTO users (username, password, real_name, student_id, login_type, role, plain_password) VALUES (?, ?, ?, ?, 'student_id', 'student', ?)",
                    (gen_username, hash_password(password), real_name or username, username, password)
                )
                imported += 1
            except Exception as e:
                errors.append(f"第{row_idx}行({username}): 导入失败 - {str(e)}")

    return JSONResponse({"success": True, "imported": imported, "errors": errors})

# ============================================================
# Admin API - Statistics
# ============================================================
@app.get("/api/admin/stats")
async def admin_stats(request: Request):
    user = require_admin(request)
    with get_db() as db:
        stats = {
            "user_count": db.execute("SELECT COUNT(*) as c FROM users WHERE role='student'").fetchone()["c"],
            "question_count": db.execute("SELECT COUNT(*) as c FROM questions").fetchone()["c"],
            "exam_count": db.execute("SELECT COUNT(*) as c FROM exams").fetchone()["c"],
            "session_count": db.execute("SELECT COUNT(*) as c FROM exam_sessions WHERE status='submitted'").fetchone()["c"],
            "avg_score": db.execute("SELECT COALESCE(AVG(total_score),0) as a FROM exam_sessions WHERE status='submitted'").fetchone()["a"],
            "pass_count": db.execute("SELECT COUNT(*) as c FROM exam_sessions es JOIN exams e ON es.exam_id=e.id WHERE es.status='submitted' AND es.total_score >= e.pass_score").fetchone()["c"],
        }

        # Question type distribution
        type_dist = db.execute("SELECT type, COUNT(*) as c FROM questions GROUP BY type").fetchall()
        stats["type_distribution"] = {r["type"]: r["c"] for r in type_dist}

    return JSONResponse(stats)

# ============================================================
# Admin API - Image Upload
# ============================================================
@app.post("/api/admin/upload/image")
async def upload_image(request: Request, file: UploadFile = File(...)):
    user = require_admin(request)
    allowed = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        return JSONResponse({"error": "仅支持 png/jpg/gif/bmp/webp 格式图片"}, status_code=400)

    filename = f"{uuid.uuid4().hex}{ext}"
    filepath = os.path.join(UPLOAD_DIR, filename)
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        return JSONResponse({"error": "图片大小不能超过10MB"}, status_code=400)
    with open(filepath, "wb") as f:
        f.write(content)

    url = f"/static/uploads/{filename}"
    return JSONResponse({"success": True, "url": url, "filename": filename})

# ============================================================
# Admin API - Word/Txt Import
# ============================================================
def _pre_normalize_text(text: str) -> str:
    """Pre-normalize pasted text: split single-line content into proper multi-line format.

    Handles the common case where users paste questions from PDF/web as a single
    continuous line with all options, answers and sub-questions run together.
    Inserts line breaks before structural markers (options, answers, sub-questions).
    """
    # Insert newline before structural markers (only when they appear mid-line)
    # Order matters: longer patterns first to avoid partial matches

    # 1. Type markers: 【单选题】, 【案例分析题】, etc.
    text = re.sub(r'(?<!\n)\s*(【[^】]*(?:题|选)】)', r'\n\1', text)
    text = re.sub(r'(?<!\n)\s*(\[[^\]]*(?:题|选)\])', r'\n\1', text)

    # 2. Answer lines: 答案： or 【答案】
    text = re.sub(r'(?<!\n)\s*(答案[：:])', r'\n\1', text)
    text = re.sub(r'(?<!\n)\s*(【答案】)', r'\n\1', text)

    # 3. Sub-question markers: 提问1：, 提问：, 提问1, 第1问
    text = re.sub(r'(?<!\n)\s*(提问\s*\d*\s*[：:]?)', r'\n\1', text)
    text = re.sub(r'(?<!\n)\s*(第\s*\d\s*问[：:]?)', r'\n\1', text)

    # 4. Option lines: A. B. C. ... H.
    # Match A-H followed by . or 、 or : but NOT inside words like "PSA" or "AMACR"
    # Use word-boundary-like check: option letter must be preceded by space or start-of-line
    text = re.sub(r'(?<=[ \t\n])([A-H])[\.、．:：]\s*', r'\n\1. ', text)
    # Also handle options at the very beginning of text (rare but possible)
    text = re.sub(r'^([A-H])[\.、．:：]\s*', r'\n\1. ', text)

    # 5. Other structural markers
    text = re.sub(r'(?<!\n)\s*(解析[：:])', r'\n\1', text)
    text = re.sub(r'(?<!\n)\s*(分值[：:])', r'\n\1', text)
    text = re.sub(r'(?<!\n)\s*(分类[：:])', r'\n\1', text)

    return text


def parse_txt_questions(text: str) -> tuple:
    """Parse questions from TXT format. Returns (questions_list, errors_list)"""
    questions = []
    errors = []

    # Pre-normalize: split single-line pasted text into proper multi-line format
    text = _pre_normalize_text(text)

    lines = text.split('\n')
    i = 0
    q_num = 0
    last_type = 'single'  # Track last seen type for auto-detection

    while i < len(lines):
        line = lines[i].strip()

        # Detect type markers - immediately start parsing the question
        if line.startswith('【单选题】') or line.startswith('[单选题]') or line.startswith('【单选】') or line == '单选题':
            last_type = 'single'
            # If next line is another type marker, just skip; otherwise start parsing
            next_is_marker = (i + 1 < len(lines) and 
                (lines[i+1].strip().startswith('【') or lines[i+1].strip().startswith('[') or
                 lines[i+1].strip() in ('单选题', '多选题', '共用题干', '案例分析题', '案例分析')))
            if not next_is_marker:
                q_num += 1
                try:
                    q, i = _parse_single_question(lines, i, 'single', q_num)
                    if q:
                        questions.append(q)
                except Exception as e:
                    errors.append(f"第{q_num}题: {str(e)}")
                    i += 1
                continue
            i += 1
            continue

        elif line.startswith('【多选题】') or line.startswith('[多选题]') or line.startswith('【多选】') or line == '多选题':
            last_type = 'multiple'
            next_is_marker = (i + 1 < len(lines) and 
                (lines[i+1].strip().startswith('【') or lines[i+1].strip().startswith('[') or
                 lines[i+1].strip() in ('单选题', '多选题', '共用题干', '案例分析题', '案例分析')))
            if not next_is_marker:
                q_num += 1
                try:
                    q, i = _parse_single_question(lines, i, 'multiple', q_num)
                    if q:
                        questions.append(q)
                except Exception as e:
                    errors.append(f"第{q_num}题: {str(e)}")
                    i += 1
                continue
            i += 1
            continue

        elif line.startswith('【共用题干】') or line.startswith('[共用题干]') or line.startswith('【共用题干单选】') or line == '共用题干':
            last_type = 'shared'
            q_num += 1
            try:
                q, i = _parse_shared_case_question(lines, i, 'shared', q_num)
                if q:
                    questions.append(q)
            except Exception as e:
                errors.append(f"第{q_num}题: {str(e)}")
                i += 1

        elif line.startswith('【案例分析题】') or line.startswith('【案例分析】') or line.startswith('[案例分析]') or line == '案例分析题' or line == '案例分析':
            last_type = 'case'
            q_num += 1
            try:
                q, i = _parse_shared_case_question(lines, i, 'case', q_num)
                if q:
                    questions.append(q)
            except Exception as e:
                errors.append(f"第{q_num}题: {str(e)}")
                i += 1

        # Auto-detect: 题目： starts a new single/multiple question (uses last_type)
        elif (line.startswith('题目：') or line.startswith('题目:')) and last_type in ('single', 'multiple'):
            q_num += 1
            try:
                q, i = _parse_single_question(lines, i, last_type, q_num)
                if q:
                    questions.append(q)
            except Exception as e:
                errors.append(f"第{q_num}题: {str(e)}")
                i += 1

        # Auto-detect: 题干： starts a new shared/case question
        elif (line.startswith('题干：') or line.startswith('题干:')) and last_type in ('shared', 'case'):
            q_num += 1
            try:
                q, i = _parse_shared_case_question(lines, i, last_type, q_num)
                if q:
                    questions.append(q)
            except Exception as e:
                errors.append(f"第{q_num}题: {str(e)}")
                i += 1

        else:
            i += 1

    return questions, errors

def _parse_single_question(lines, start, q_type, q_num):
    """Parse a single/multiple choice question starting from the type marker line."""
    i = start
    # If start line is a type marker, skip it; otherwise it has content (auto-detect)
    line0 = lines[start].strip()
    if line0.startswith('\u3010') or line0.startswith('['):
        i = start + 1
    content_lines = []
    options = []
    answer = ""
    explanation = ""
    score = 1
    category = ""

    # Read content until first option
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith('分类') or line.startswith('【分类】'):
            cat_match = re.search(r'[：:]\s*(.+)', line)
            if cat_match:
                category = cat_match.group(1).strip()
            i += 1
            continue
        if re.match(r'^[A-H][\.、．:：\s]', line):
            break
        if line.startswith('答案') or line.startswith('【答案】'):
            break
        if line.startswith('解析') or line.startswith('【解析】'):
            break
        if line.startswith('分值') or line.startswith('【分值】'):
            break
        if re.match(r'^【', line):
            break
        content_lines.append(line)
        i += 1

    content = '\n'.join(content_lines).strip()
    # Strip common prefixes
    content = re.sub(r'^题目[：:]\s*', '', content).strip()

    # Read options
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        m = re.match(r'^([A-H])[\.、．:：\s]+(.+)$', line)
        if m:
            options.append({"label": m.group(1), "text": m.group(2).strip()})
            i += 1
        else:
            break

    # Read answer, explanation, score
    found_answer = False
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            if found_answer:
                break
            i += 1
            continue

        # Stop at new question boundary
        if line.startswith('题目：') or line.startswith('题目:'):
            break
        if line.startswith('题干：') or line.startswith('题干:'):
            break
        if re.match(r'^【[单多共案]', line):
            break

        if line.startswith('答案') or line.startswith('【答案】'):
            ans_match = re.search(r'[：:]\s*(.+)', line)
            if ans_match:
                answer = re.sub(r'[（(]\s*\d+\s*分\s*[）)]', '', ans_match.group(1).strip()).strip()
                found_answer = True
            else:
                # Handle 【答案】C format (no colon)
                ans_text = re.sub(r'^[【\[]答案[】\]]\s*', '', line).strip()
                ans_text = re.sub(r'[（(]\s*\d+\s*分\s*[）)]', '', ans_text).strip()
                if ans_text:
                    answer = ans_text
                    found_answer = True
            i += 1
        elif line.startswith('解析') or line.startswith('【解析】'):
            exp_match = re.search(r'[：:]\s*(.+)', line)
            if exp_match:
                explanation = exp_match.group(1).strip()
            i += 1
        elif line.startswith('分值') or line.startswith('【分值】'):
            score_match = re.search(r'[：:]\s*(.+)', line)
            if score_match:
                try:
                    score = float(score_match.group(1).strip())
                except:
                    pass
            i += 1
        else:
            i += 1

    if not content:
        return None, i

    return {
        "type": q_type,
        "content": content,
        "options": options,
        "answer": answer,
        "explanation": explanation,
        "score": score,
        "sub_questions": [],
        "category": category
    }, i

def _split_trailing_stem(content):
    """Split content into (main_stem, sub_stem) if content ends with a sub-question stem.

    Heuristic: if the last sentence/phrase of content doesn't end with sentence-ending
    punctuation (。！？；) and the text after the last such punctuation looks like a
    question stem (short phrase), split it out as the first sub-question's stem.
    """
    last_end = -1
    for ch in ['。', '！', '？', '；']:
        idx = content.rfind(ch)
        if idx > last_end:
            last_end = idx

    if last_end < 0 or last_end >= len(content) - 1:
        last_nl = content.rfind('\n')
        if last_nl > 0 and last_nl < len(content) - 1:
            main = content[:last_nl].strip()
            sub = content[last_nl+1:].strip()
            if len(sub) < 60 and sub:
                return main, sub
        return content, ""

    main = content[:last_end+1].strip()
    sub = content[last_end+1:].strip()

    if not sub or len(sub) > 80:
        return content, ""

    return main, sub


def _parse_shared_case_question(lines, start, q_type, q_num):
    """Parse a shared/case question with sub-questions.

    Supports multiple sub-question identification patterns:
    1. Explicit markers: 提问1, 第1问, 1., 【提问】
    2. Auto-detection: after "答案：" line, any non-option/non-marker text
       is treated as a new sub-question stem (handles formats without
       explicit sub-question numbering).
    """
    i = start
    line0 = lines[start].strip()
    if line0.startswith('\u3010') or line0.startswith('['):
        i = start + 1
    content_lines = []
    sub_questions = []
    explanation = ""
    score = 3 if q_type == 'case' else 1
    category = ""

    # Read main content (题干) until sub-question markers or first option line
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue
        if line.startswith('分类') or line.startswith('【分类】'):
            cat_match = re.search(r'[：:]\s*(.+)', line)
            if cat_match:
                category = cat_match.group(1).strip()
            i += 1
            continue
        if line.startswith('【提问】') or line.startswith('[提问]') or line.startswith('【子题】'):
            i += 1
            break
        if re.match(r'^提问\s*\d*\s*[：:]', line) or re.match(r'^提问\s+\d', line) or re.match(r'^第\s*\d\s*问', line):
            break
        if line.startswith('解析') or line.startswith('【解析】'):
            exp_match = re.search(r'[：:]\s*(.+)', line)
            if exp_match:
                explanation = exp_match.group(1).strip()
            i += 1
            continue
        if line.startswith('分值') or line.startswith('【分值】'):
            score_match = re.search(r'[：:]\s*(.+)', line)
            if score_match:
                try:
                    score = float(score_match.group(1).strip())
                except:
                    pass
            i += 1
            continue
        if re.match(r'^【', line) and not line.startswith('【提问】'):
            break
        # Skip bare type markers (案例分析题, 共用题干, etc.) that may slip in from preprocessing
        if line in ('案例分析题', '案例分析', '共用题干', '共用题干单选', '单选题', '多选题'):
            i += 1
            continue
        # If we hit an option line (A-H.), the sub-question zone has started
        if re.match(r'^([A-H])[\.、．:：\s]+', line):
            break
        content_lines.append(line)
        i += 1

    content = '\n'.join(content_lines).strip()

    # Helper: check if a line is a known structural marker that is NOT sub-question content
    def _is_boundary_marker(line):
        """Check if line is a top-level boundary (new question type marker)"""
        if re.match(r'^【(单选题|多选题|共用题干|案例分析|单选|多选)', line):
            return True
        if line in ('案例分析题', '案例分析', '共用题干', '共用题干单选', '单选题', '多选题'):
            return True
        if line.startswith('题目：') or line.startswith('题目:'):
            return True
        if line.startswith('题干：') or line.startswith('题干:'):
            return True
        return False

    # Helper: check if a line looks like a new option
    def _is_option_line(line):
        return bool(re.match(r'^([A-H])[\.、．:：\s]+(.+)$', line))

    # Parse sub-questions
    current_sub = None
    sub_has_answer = False  # Track whether current sub already has an answer filled

    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        # Stop at boundary markers (new question type)
        if _is_boundary_marker(line):
            break

        # 1) Explicit sub-question markers: 提问：, 提问1：, 提问1, 第1问
        if re.match(r'^提问\s*\d*\s*[：:]', line) or re.match(r'^提问\s+\d', line) or re.match(r'^第\s*\d\s*问', line):
            if current_sub:
                sub_questions.append(current_sub)
            sub_content = re.sub(r'^提问\s*\d*\s*[：:]\s*|^提问\s+\d+\s*[：:]?\s*|^第\s*\d\s*问\s*[：:]?\s*', '', line).strip()
            current_sub = {"content": sub_content, "options": [], "answer": ""}
            sub_has_answer = False
            i += 1
            continue

        # 2) Numbered sub-question: "1." "2." etc. (only if we already have a sub or are in sub-question zone)
        if re.match(r'^\d+[\.、．:：]\s*', line) and (current_sub is not None or len(content_lines) > 0):
            if current_sub:
                sub_questions.append(current_sub)
            sub_content = re.sub(r'^\d+[\.、．:：]\s*', '', line).strip()
            current_sub = {"content": sub_content, "options": [], "answer": ""}
            sub_has_answer = False
            i += 1
            continue

        # 3) Option line A-H
        if _is_option_line(line):
            # If no current sub exists yet, create one with empty content
            # (the stem text may have been absorbed into the main content)
            if current_sub is None:
                current_sub = {"content": "", "options": [], "answer": ""}
                sub_has_answer = False
            m = re.match(r'^([A-H])[\.、．:：\s]+(.+)$', line)
            if m:
                current_sub["options"].append({"label": m.group(1), "text": m.group(2).strip()})
            i += 1
            continue

        # 4) Answer line
        if line.startswith('答案') or line.startswith('【答案】'):
            if current_sub is None:
                current_sub = {"content": "", "options": [], "answer": ""}
            ans_match = re.search(r'[：:]\s*(.+)', line)
            if ans_match:
                raw_ans = ans_match.group(1).strip()
            else:
                # Handle 【答案】C format (no colon)
                raw_ans = re.sub(r'^[【\[]答案[】\]]\s*', '', line).strip()
            # Clean up answer: remove score annotations like (1分),（2分）, etc.
            raw_ans = re.sub(r'[（(]\s*\d+\s*分\s*[）)]', '', raw_ans).strip()
            if raw_ans:
                current_sub["answer"] = raw_ans
            sub_has_answer = True
            i += 1
            continue

        # 5) Explanation line (within sub-question area - belongs to overall question)
        if line.startswith('解析') or line.startswith('【解析】'):
            exp_match = re.search(r'[：:]\s*(.+)', line)
            if exp_match:
                explanation = exp_match.group(1).strip() if not explanation else explanation + '\n' + exp_match.group(1).strip()
            i += 1
            continue

        # 6) Score line
        if line.startswith('分值') or line.startswith('【分值】'):
            score_match = re.search(r'[：:]\s*(.+)', line)
            if score_match:
                try:
                    score = float(score_match.group(1).strip())
                except:
                    pass
            i += 1
            continue

        # 7) Auto-detect new sub-question:
        #    If current sub already has an answer filled, and we encounter
        #    a line that is not an option/answer/marker, it's a new sub-question stem.
        #    This handles formats like:
        #        A. 选项1
        #        B. 选项2
        #        答案：B
        #        下一个子题的题干文字...   <-- this triggers auto-detection
        #        A. 选项1
        #        ...
        if sub_has_answer and current_sub is not None:
            # Current sub is complete (has answer), save it
            sub_questions.append(current_sub)
            # This line is the stem of a new sub-question
            current_sub = {"content": line, "options": [], "answer": ""}
            sub_has_answer = False
            i += 1
            continue

        # 8) If we have a current sub but no answer yet, this line might be:
        #    - continuation of sub-question content (multi-line stem)
        #    - or a stray line we can't classify
        if current_sub is not None and not sub_has_answer:
            # Append to current sub's content if it's not an option/answer
            if current_sub["content"]:
                current_sub["content"] += '\n' + line
            else:
                current_sub["content"] = line
            i += 1
            continue

        # 9) No current sub and no auto-detect: skip
        i += 1

    if current_sub:
        sub_questions.append(current_sub)

    # Post-processing: if first sub-question has empty stem,
    # try to extract it from the tail of main content
    if sub_questions and not sub_questions[0]["content"]:
        main_stem, sub_stem = _split_trailing_stem(content)
        if sub_stem:
            content = main_stem
            sub_questions[0]["content"] = sub_stem

    if not content:
        return None, i

    return {
        "type": q_type,
        "content": content,
        "options": [],
        "answer": "",
        "explanation": explanation,
        "score": score,
        "sub_questions": sub_questions,
        "category": category
    }, i



@app.post("/api/admin/import/preview")
async def import_preview(request: Request):
    """Preview parsed questions from pasted text without saving to DB"""
    user = require_admin(request)
    body = await request.json()
    text = body.get("text", "")
    if not text.strip():
        return JSONResponse({"questions": [], "errors": ["请输入试题内容"]})

    questions, errors = parse_txt_questions(text)

    # Format for display
    result = []
    for q in questions:
        item = {
            "type": q["type"],
            "type_label": {"single": "单选题", "multiple": "多选题", "shared": "共用题干", "case": "案例分析"}.get(q["type"], q["type"]),
            "content": q["content"][:100] + ("..." if len(q["content"]) > 100 else ""),
            "options_count": len(q.get("options", [])),
            "answer": q.get("answer", ""),
            "explanation": q.get("explanation", "")[:80],
            "category": q.get("category", ""),
            "sub_count": len(q.get("sub_questions", [])),
            "sub_questions": []
        }
        for sq in q.get("sub_questions", []):
            item["sub_questions"].append({
                "content": sq.get("content", "")[:60],
                "answer": sq.get("answer", ""),
                "options_count": len(sq.get("options", []))
            })
        result.append(item)

    return JSONResponse({"questions": result, "errors": errors, "total": len(result)})


def parse_docx_questions(content: bytes) -> tuple:
    """Parse questions from Word (.docx) format. Returns (questions_list, errors_list)"""
    try:
        from docx import Document
    except ImportError:
        return [], ["python-docx 库未安装，请在服务器执行: source /opt/exam_system/venv/bin/activate && pip install python-docx"]

    questions = []
    errors = []

    try:
        doc = Document(BytesIO(content))
    except Exception as e:
        return [], [f"Word文件解析失败: {str(e)}"]

    # Extract images from document
    image_map = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            try:
                img_data = rel.target_part.blob
                ext = os.path.splitext(rel.target_part.partname)[1].lower()
                if ext not in {'.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp'}:
                    ext = '.png'
                filename = f"{uuid.uuid4().hex}{ext}"
                filepath = os.path.join(UPLOAD_DIR, filename)
                with open(filepath, "wb") as f:
                    f.write(img_data)
                image_map[rel.target_part.partname] = f"/static/uploads/{filename}"
            except Exception:
                pass

    # Build text with image placeholders
    full_text_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        # Check for inline images in runs
        for run in para.runs:
            if run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                for drawing in drawings:
                    blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                    for blip in blips:
                        embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed and embed in doc.part.rels:
                            img_rel = doc.part.rels[embed]
                            partname = img_rel.target_part.partname
                            if partname in image_map:
                                text += f"\n[图片:{image_map[partname]}]\n"

        if text:
            full_text_parts.append(text)

    full_text = '\n'.join(full_text_parts)
    # Now parse the text like TXT format
    questions, errors = parse_txt_questions(full_text)

    # Convert image placeholders to <img> tags in content
    for q in questions:
        q["content"] = re.sub(r'\[图片:(/static/uploads/[^\]]+)\]', r'<img src="\1" style="max-width:100%;border-radius:8px;margin:8px 0;">', q["content"])
        for sub in q.get("sub_questions", []):
            sub["content"] = re.sub(r'\[图片:(/static/uploads/[^\]]+)\]', r'<img src="\1" style="max-width:100%;border-radius:8px;margin:8px 0;">', sub.get("content", ""))

    return questions, errors



def auto_create_exam_from_source_set(db, source_set: str, question_ids: list):
    """If source_set is provided, auto-create an exam with all imported questions."""
    if not source_set or not question_ids:
        return None

    # Check if exam with same title already exists
    existing = db.execute("SELECT id FROM exams WHERE title=?", (source_set,)).fetchone()
    if existing:
        exam_id = existing["id"]
        # Append new questions to existing exam
        max_order = db.execute("SELECT MAX(display_order) FROM exam_questions WHERE exam_id=?", (exam_id,)).fetchone()[0] or 0
        for idx, qid in enumerate(question_ids):
            # Check if question already in exam
            already = db.execute("SELECT id FROM exam_questions WHERE exam_id=? AND question_id=?", (exam_id, qid)).fetchone()
            if not already:
                max_order += 1
                db.execute("INSERT INTO exam_questions (exam_id, question_id, display_order, score_override) VALUES (?, ?, ?, ?)",
                           (exam_id, qid, max_order, None))
        # Recalculate total_score
        total = db.execute("""
            SELECT COALESCE(SUM(COALESCE(eq.score_override, q.score)), 0)
            FROM exam_questions eq JOIN questions q ON eq.question_id = q.id
            WHERE eq.exam_id = ?
        """, (exam_id,)).fetchone()[0]
        db.execute("UPDATE exams SET total_score=? WHERE id=?", (total, exam_id))
    else:
        # Create new exam
        total_score = 0
        for qid in question_ids:
            q = db.execute("SELECT score FROM questions WHERE id=?", (qid,)).fetchone()
            if q:
                total_score += q["score"]

        cursor = db.execute(
            """INSERT INTO exams (title, description, duration, total_score, pass_score, is_practice, shuffle_questions, status)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?)""",
            (source_set, f"由套卷\u300c{source_set}\u300d自动生成的考试", 120, total_score, total_score * 0.6, False, False, "active")
        )
        exam_id = cursor.lastrowid

        for idx, qid in enumerate(question_ids):
            db.execute("INSERT INTO exam_questions (exam_id, question_id, display_order, score_override) VALUES (?, ?, ?, ?)",
                       (exam_id, qid, idx + 1, None))

    return exam_id

@app.post("/api/admin/import/txt")
async def import_txt(request: Request, file: UploadFile = File(...), source_set: str = Form("")):
    user = require_admin(request)
    content = await file.read()

    try:
        text = content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            text = content.decode("gbk")
        except:
            return JSONResponse({"error": "文件编码错误，请使用UTF-8或GBK编码"}, status_code=400)

    questions, errors = parse_txt_questions(text)

    count = 0
    question_ids = []
    with get_db() as db:
        for q in questions:
            try:
                cursor = db.execute(
                    """INSERT INTO questions (type, category, content, options, answer, explanation, score, sub_questions, images, source_set)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (q["type"], q.get("category", ""), q["content"],
                     json.dumps(q.get("options", []), ensure_ascii=False),
                     q.get("answer", ""), q.get("explanation", ""),
                     q.get("score", 1),
                     json.dumps(q.get("sub_questions", []), ensure_ascii=False),
                     "[]",
                     source_set)
                )
                question_ids.append(cursor.lastrowid)
                count += 1
            except Exception as e:
                errors.append(f"导入失败: {str(e)}")

        # Auto-create exam if source_set is provided
        exam_id = auto_create_exam_from_source_set(db, source_set, question_ids)

    result = {"success": True, "imported": count, "errors": errors, "source_set": source_set}
    if exam_id:
        result["exam_id"] = exam_id
        result["message"] = f"已自动创建考试「{source_set}」"
    return JSONResponse(result)


@app.post("/api/admin/import/word")
async def import_word(request: Request, file: UploadFile = File(...), source_set: str = Form("")):
    user = require_admin(request)
    content = await file.read()

    questions, errors = parse_docx_questions(content)

    count = 0
    question_ids = []
    with get_db() as db:
        for q in questions:
            try:
                # Extract image URLs from content
                img_urls = re.findall(r'/static/uploads/[a-f0-9]+\.\w+', q["content"])
                cursor = db.execute(
                    """INSERT INTO questions (type, category, content, options, answer, explanation, score, sub_questions, images, source_set)
                       VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                    (q["type"], q.get("category", ""), q["content"],
                     json.dumps(q.get("options", []), ensure_ascii=False),
                     q.get("answer", ""), q.get("explanation", ""),
                     q.get("score", 1),
                     json.dumps(q.get("sub_questions", []), ensure_ascii=False),
                     json.dumps(img_urls),
                     source_set)
                )
                question_ids.append(cursor.lastrowid)
                count += 1
            except Exception as e:
                errors.append(f"导入失败: {str(e)}")

        # Auto-create exam if source_set is provided
        exam_id = auto_create_exam_from_source_set(db, source_set, question_ids)

    result = {"success": True, "imported": count, "errors": errors, "source_set": source_set}
    if exam_id:
        result["exam_id"] = exam_id
        result["message"] = f"已自动创建考试「{source_set}」"
    return JSONResponse(result)

@app.get("/api/admin/import/format-guide")
async def import_format_guide(request: Request):
    """Return the format specification for Word/TXT import"""
    user = require_admin(request)
    guide = """═══════════════════════════════════════
  试题导入格式规范 (Word / TXT 通用)
═══════════════════════════════════════

支持的文件格式: .docx (Word), .txt (纯文本UTF-8/GBK)
Word文件中嵌入的图片会自动提取并保存。

───────────────────────────────────────
一、单项选择题（选项6-8个）
───────────────────────────────────────
【单选题】
分类：总论-细胞适应与损伤
关于细胞凋亡的描述，错误的是
A. 是由基因控制的程序性死亡
B. 可引起炎症反应
C. 是单个细胞的死亡
D. 不引起炎症反应
E. 可见于生理和病理过程
F. 仅发生于胚胎发育阶段
G. 又称细胞程序性死亡
答案：B
解析：细胞凋亡一般不引起炎症反应，这是与坏死的重要区别。
分值：1

───────────────────────────────────────
二、多项选择题（选项6-8个）
───────────────────────────────────────
【多选题】
分类：总论-肿瘤
下列属于癌前病变的有
A. 黏膜白斑
B. 慢性宫颈炎伴宫颈糜烂
C. 纤维囊性乳腺病
D. 肝血管瘤
E. 结肠腺瘤性息肉
F. 慢性萎缩性胃炎
G. 皮肤慢性溃疡
答案：ABCEFG
解析：肝血管瘤是良性肿瘤，不属于癌前病变。
分值：1

───────────────────────────────────────
三、共用题干单选题（单选，一个题干下设3-6小题，每题选项6-8个）
───────────────────────────────────────
【共用题干】
分类：各论-呼吸及胸腔系统
患者男，65岁，吸烟40年，持续咳嗽、咳痰3个月，近1周痰中带血。
查体：体温37.2℃，右下肺可闻及局限性哮鸣音。胸部X线示右下肺门区团块影。
提问1：最可能的诊断是
A. 肺结核
B. 肺癌
C. 慢性支气管炎
D. 肺脓肿
E. 肺炎
F. 支气管扩张
答案：B
提问2：为明确诊断，首选检查是
A. 胸部X线
B. 痰脱落细胞学
C. 支气管镜
D. CT
E. MRI
F. 肿瘤标志物
答案：C
提问3：该患者最可能的病理类型是
A. 鳞状细胞癌
B. 腺癌
C. 小细胞癌
D. 大细胞癌
E. 细支气管肺泡癌
F. 类癌
答案：A
分值：1

───────────────────────────────────────
四、案例分析题（不定项选择题，选项6-8个，下设6-8小题）
───────────────────────────────────────
【案例分析】
分类：各论-女性生殖系统
患者女，45岁，发现右乳肿块1周。查体：右乳外上象限可及一3cm×2cm肿块，
质硬，边界不清，活动度差，右乳头内陷，右腋窝可触及2枚肿大淋巴结。
提问1：该患者最可能的诊断
A. 乳腺纤维腺瘤
B. 乳腺癌
C. 乳腺囊性增生
D. 乳腺炎
E. 乳腺导管内乳头状瘤
F. 乳房脂肪坏死
答案：B
提问2：确诊依据包括
A. 肿块质硬固定
B. 乳头内陷
C. 腋窝淋巴结肿大
D. 年龄
E. 肿块边界不清
F. 活动度差
答案：ABCEF
提问3：为明确诊断应进行的检查
A. 乳腺B超
B. 乳腺钼靶X线
C. 穿刺活检
D. MRI
E. CT
F. 肿瘤标志物
答案：ABC
提问4：若确诊为乳腺癌，最常见的病理类型为
A. 浸润性导管癌
B. 浸润性小叶癌
C. 髓样癌
D. 黏液癌
E. 导管内癌
F. 炎性乳腺癌
答案：A
提问5：该患者腋窝淋巴结的处理原则
A. 行腋窝淋巴结清扫
B. 仅行前哨淋巴结活检
C. 术后放疗
D. 化疗即可
E. 不需处理
F. 视术中情况决定
答案：A
提问6：术后辅助化疗的常用方案
A. AC方案（阿霉素+环磷酰胺）
B. TC方案（多西他赛+环磷酰胺）
C. CMF方案
D. 紫杉醇单药
E. 内分泌治疗
F. 靶向治疗
答案：ABC
分值：3

═══════════════════════════════════════
注意事项：
1. 每道题目之间用空行分隔
2. 选项标签支持 A-H（最多8个选项）
3. 共用题干题：每个提问只选一个答案（单选）
4. 案例分析题：每个提问可选多个答案（不定项选择）
5. 分类为可选项，格式为"分类：xxx"，放在题型标记之后
   总论分类：细胞适应与损伤、损伤修复、血液循环障碍、炎症、肿瘤、免疫病理、遗传性疾病、其他
   各论分类：心血管系统、呼吸及胸腔系统、消化系统、泌尿和男性生殖系统、淋巴造血系统、
            内分泌和神经内分泌系统、神经系统、女性生殖系统、骨与关节、皮肤与软组织、
            头颈部、传染病与寄生虫、病理学技术、其他
6. 解析和分值为可选项
═══════════════════════════════════════"""
    return JSONResponse({"guide": guide})

# ============================================================
# Admin API - Question Source Sets (套卷管理)
# ============================================================
@app.get("/api/admin/question-sets")
async def list_question_sets(request: Request):
    """List all question source sets with statistics"""
    user = require_admin(request)
    with get_db() as db:
        sets = db.execute("""
            SELECT source_set,
                   COUNT(*) as total_count,
                   SUM(CASE WHEN type='single' THEN 1 ELSE 0 END) as single_count,
                   SUM(CASE WHEN type='multiple' THEN 1 ELSE 0 END) as multiple_count,
                   SUM(CASE WHEN type='shared' THEN 1 ELSE 0 END) as shared_count,
                   SUM(CASE WHEN type='case' THEN 1 ELSE 0 END) as case_count
            FROM questions
            WHERE source_set != '' AND source_set IS NOT NULL
            GROUP BY source_set
            ORDER BY MAX(created_at) DESC
        """).fetchall()
    return JSONResponse({"items": [dict(s) for s in sets]})

@app.get("/api/admin/questions/by-set")
async def get_questions_by_set(request: Request, source_set: str):
    """Get all questions from a specific source set"""
    user = require_admin(request)
    with get_db() as db:
        questions = db.execute(
            "SELECT * FROM questions WHERE source_set=? ORDER BY type, id",
            (source_set,)
        ).fetchall()
    return JSONResponse({"items": [dict(q) for q in questions]})

# ============================================================
# Admin API - Exam Status Toggle (Publish/Unpublish)
# ============================================================
@app.put("/api/admin/exams/{eid}/status")
async def toggle_exam_status(eid: int, request: Request):
    user = require_admin(request)
    data = await request.json()
    new_status = data.get("status", "")
    if new_status not in ("active", "draft"):
        return JSONResponse({"error": "无效状态，仅支持 active(发布) 和 draft(下架)"}, status_code=400)

    with get_db() as db:
        db.execute("UPDATE exams SET status=? WHERE id=?", (new_status, eid))
    return JSONResponse({"success": True, "status": new_status})

# ============================================================
# Admin API - Random Question Selection
# ============================================================
@app.post("/api/admin/questions/random")
async def get_random_questions(request: Request):
    """Random question selection with custom counts per type.
    Accepts JSON body with type_counts: {single: n, multiple: n, shared: n, case: n}
    Or legacy GET-style: q_type, count, exclude_ids
    """
    user = require_admin(request)
    data = await request.json()

    # Support both new format (type_counts) and legacy format (q_type/count)
    type_counts = data.get("type_counts", {})  # {single: 10, multiple: 5, ...}
    exclude_ids = data.get("exclude_ids", "")

    # Legacy support
    if not type_counts:
        q_type = data.get("q_type") or data.get("type")
        count = data.get("count", 10)
        if q_type:
            type_counts = {q_type: count}
        else:
            type_counts = {"single": count}  # default to single choice

    with get_db() as db:
        selected_questions = []
        errors = []
        total_available = {}

        # First, get available counts for each type
        for q_type in ["single", "multiple", "shared", "case"]:
            available = db.execute(
                "SELECT COUNT(*) as c FROM questions WHERE type=?", (q_type,)
            ).fetchone()["c"]
            total_available[q_type] = available

        # Then, select questions for each type
        exclude_list = [x.strip() for x in exclude_ids.split(",") if x.strip()] if exclude_ids else []

        for q_type, count in type_counts.items():
            if q_type not in ["single", "multiple", "shared", "case"]:
                continue
            if count <= 0:
                continue

            # Check availability
            available = total_available.get(q_type, 0)
            if available < count:
                type_names = {
                    "single": "单选题",
                    "multiple": "多选题",
                    "shared": "共用题干题",
                    "case": "案例分析题"
                }
                unit = "组" if q_type in ("shared", "case") else "道"
                errors.append(f"{type_names.get(q_type, q_type)}：需要 {count} {unit}，仅有 {available} {unit}")
                continue

            # Build query
            query = "SELECT * FROM questions WHERE type=?"
            params = [q_type]

            if exclude_list:
                placeholders = ",".join(["?"] * len(exclude_list))
                query += f" AND id NOT IN ({placeholders})"
                params.extend(exclude_list)

            query += " ORDER BY RANDOM() LIMIT ?"
            params.append(count)

            questions = db.execute(query, params).fetchall()

            # If not enough after excluding, try without exclusion
            if len(questions) < count:
                questions = db.execute(
                    "SELECT * FROM questions WHERE type=? ORDER BY RANDOM() LIMIT ?",
                    (q_type, count)
                ).fetchall()

            selected_questions.extend([dict(q) for q in questions])

        if errors:
            return JSONResponse({
                "error": "抽题失败：\n" + "\n".join(errors),
                "available": total_available,
                "selected": []
            })

    return JSONResponse({
        "questions": selected_questions,
        "available": total_available,
        "type_counts": type_counts
    })


# ============================================================
# Blog API
# ============================================================

@app.get("/api/blog/posts")
async def list_blog_posts(request: Request, page: int = 1, size: int = 10, category: str = "", published: str = ""):
    """List blog posts - admin sees all, students see only published"""
    user = get_current_user(request)
    
    conditions = []
    params = []
    
    # Unauthenticated or non-admin users only see published posts
    if not user or user["role"] != "admin":
        conditions.append("is_published = 1")
    elif published != "":
        if published == "1":
            conditions.append("is_published = 1")
        elif published == "0":
            conditions.append("is_published = 0")
    
    if category:
        conditions.append("category = ?")
        params.append(category)
    
    where = " WHERE " + " AND ".join(conditions) if conditions else ""
    
    with get_db() as db:
        total = db.execute(f"SELECT COUNT(*) FROM blog_posts{where}", params).fetchone()[0]
        
        offset = (page - 1) * size
        rows = db.execute(
            f"SELECT * FROM blog_posts{where} ORDER BY is_pinned DESC, created_at DESC LIMIT ? OFFSET ?",
            params + [size, offset]
        ).fetchall()
    
    posts = [dict(r) for r in rows]
    return JSONResponse({"posts": posts, "total": total, "page": page, "size": size})


@app.get("/api/blog/posts/{post_id}")
async def get_blog_post(post_id: int, request: Request):
    """Get a single blog post"""
    user = get_current_user(request)
    
    with get_db() as db:
        post = db.execute("SELECT * FROM blog_posts WHERE id=?", (post_id,)).fetchone()
        if not post:
            return JSONResponse({"error": "文章不存在"}, status_code=404)
        
        post_dict = dict(post)
        
        # Unauthenticated or non-admin users can only view published posts
        if (not user or user["role"] != "admin") and not post_dict["is_published"]:
            return JSONResponse({"error": "文章未发布"}, status_code=403)
        
        # Increment view count
        db.execute("UPDATE blog_posts SET view_count = view_count + 1 WHERE id=?", (post_id,))
    
    return JSONResponse(post_dict)


@app.post("/api/blog/posts")
async def create_blog_post(request: Request):
    """Create a new blog post (admin only)"""
    user = get_current_user(request)
    if user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)
    
    data = await request.json()
    title = data.get("title", "").strip()
    if not title:
        return JSONResponse({"error": "标题不能为空"}, status_code=400)
    
    post_content = data.get("content", "")
    summary = data.get("summary", "")
    category = data.get("category", "学习笔记")
    tags = data.get("tags", "")
    cover_image = data.get("cover_image", "")
    is_published = 1 if data.get("is_published") else 0
    is_pinned = 1 if data.get("is_pinned") else 0
    author = data.get("author", "张老师")
    
    # Generate slug from title
    import hashlib
    slug = hashlib.md5(f"{title}{__import__('time').time()}".encode()).hexdigest()[:12]
    
    with get_db() as db:
        cursor = db.execute(
            """INSERT INTO blog_posts (title, slug, content, summary, cover_image, category, tags, author, is_published, is_pinned)
               VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)""",
            (title, slug, post_content, summary, cover_image, category, tags, author, is_published, is_pinned)
        )
    
    return JSONResponse({"id": cursor.lastrowid, "title": title, "slug": slug})


@app.put("/api/blog/posts/{post_id}")
async def update_blog_post(post_id: int, request: Request):
    """Update a blog post (admin only)"""
    user = get_current_user(request)
    if user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)
    
    data = await request.json()
    
    with get_db() as db:
        post = db.execute("SELECT * FROM blog_posts WHERE id=?", (post_id,)).fetchone()
        if not post:
            return JSONResponse({"error": "文章不存在"}, status_code=404)
        
        title = data.get("title", post["title"])
        post_content = data.get("content", post["content"])
        summary = data.get("summary", post["summary"])
        category = data.get("category", post["category"])
        tags = data.get("tags", post["tags"])
        cover_image = data.get("cover_image", post["cover_image"])
        author = data.get("author", post["author"])
        is_published = 1 if data.get("is_published") else 0
        is_pinned = 1 if data.get("is_pinned") else 0
        
        db.execute(
            """UPDATE blog_posts SET title=?, content=?, summary=?, cover_image=?, category=?, tags=?, author=?, is_published=?, is_pinned=?, updated_at=CURRENT_TIMESTAMP
               WHERE id=?""",
            (title, post_content, summary, cover_image, category, tags, author, is_published, is_pinned, post_id)
        )
    
    return JSONResponse({"id": post_id, "title": title})


@app.delete("/api/blog/posts/{post_id}")
async def delete_blog_post(post_id: int, request: Request):
    """Delete a blog post (admin only)"""
    user = get_current_user(request)
    if user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)
    
    with get_db() as db:
        db.execute("DELETE FROM blog_posts WHERE id=?", (post_id,))
    
    return JSONResponse({"success": True})


@app.get("/api/blog/categories")
async def list_blog_categories(request: Request):
    """List all blog categories with post counts"""
    user = get_current_user(request)
    
    with get_db() as db:
        if user and user["role"] == "admin":
            rows = db.execute(
                "SELECT category, COUNT(*) as count FROM blog_posts GROUP BY category ORDER BY count DESC"
            ).fetchall()
        else:
            rows = db.execute(
                "SELECT category, COUNT(*) as count FROM blog_posts WHERE is_published=1 GROUP BY category ORDER BY count DESC"
            ).fetchall()
    
    categories = [{"name": r["category"], "count": r["count"]} for r in rows]
    return JSONResponse({"categories": categories})


# ============================================================
# Ads API - 广告管理
# ============================================================

@app.get("/api/ads")
async def list_active_ads(request: Request):
    """获取所有启用的广告（前端展示用）"""
    with get_db() as db:
        ads = db.execute(
            "SELECT id, title, image_url, link_url FROM ads WHERE is_active=1 ORDER BY sort_order ASC, id ASC"
        ).fetchall()
    return JSONResponse([dict(a) for a in ads])

@app.get("/api/admin/ads")
async def admin_list_ads(request: Request):
    """管理员获取所有广告"""
    user = require_admin(request)
    with get_db() as db:
        ads = db.execute("SELECT * FROM ads ORDER BY sort_order ASC, id ASC").fetchall()
    return JSONResponse([dict(a) for a in ads])

@app.post("/api/admin/ads")
async def create_ad(request: Request):
    """新增广告（admin only）"""
    user = require_admin(request)
    data = await request.json()
    title = data.get("title", "")
    image_url = data.get("image_url", "")
    link_url = data.get("link_url", "")
    is_active = data.get("is_active", 1)
    sort_order = data.get("sort_order", 0)
    if not image_url:
        return JSONResponse({"error": "图片不能为空"}, status_code=400)
    with get_db() as db:
        db.execute(
            "INSERT INTO ads (title, image_url, link_url, is_active, sort_order) VALUES (?, ?, ?, ?, ?)",
            (title, image_url, link_url, is_active, sort_order)
        )
        ad = db.execute("SELECT * FROM ads WHERE id=last_insert_rowid()").fetchone()
    return JSONResponse(dict(ad))

@app.put("/api/admin/ads/{ad_id}")
async def update_ad(ad_id: int, request: Request):
    """更新广告（admin only）"""
    user = require_admin(request)
    data = await request.json()
    with get_db() as db:
        ad = db.execute("SELECT * FROM ads WHERE id=?", (ad_id,)).fetchone()
        if not ad:
            return JSONResponse({"error": "广告不存在"}, status_code=404)
        title = data.get("title", ad["title"])
        image_url = data.get("image_url", ad["image_url"])
        link_url = data.get("link_url", ad["link_url"])
        is_active = data.get("is_active", ad["is_active"])
        sort_order = data.get("sort_order", ad["sort_order"])
        db.execute(
            "UPDATE ads SET title=?, image_url=?, link_url=?, is_active=?, sort_order=? WHERE id=?",
            (title, image_url, link_url, is_active, sort_order, ad_id)
        )
        ad = db.execute("SELECT * FROM ads WHERE id=?", (ad_id,)).fetchone()
    return JSONResponse(dict(ad))

@app.delete("/api/admin/ads/{ad_id}")
async def delete_ad(ad_id: int, request: Request):
    """删除广告（admin only）"""
    user = require_admin(request)
    with get_db() as db:
        db.execute("DELETE FROM ads WHERE id=?", (ad_id,))
    return JSONResponse({"success": True})

@app.post("/api/admin/ads/upload")
async def upload_ad_image(request: Request, file: UploadFile = File(...)):
    """上传广告图片（admin only）"""
    user = require_admin(request)
    allowed = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
    ext = os.path.splitext(file.filename)[1].lower()
    if ext not in allowed:
        return JSONResponse({"error": "仅支持 png/jpg/gif/bmp/webp 格式图片"}, status_code=400)
    ads_dir = os.path.join(UPLOAD_DIR, "ads")
    os.makedirs(ads_dir, exist_ok=True)
    filename = f"{uuid.uuid4().hex}{ext}"
    filepath = os.path.join(ads_dir, filename)
    content = await file.read()
    if len(content) > 10 * 1024 * 1024:
        return JSONResponse({"error": "图片大小不能超过10MB"}, status_code=400)
    with open(filepath, "wb") as f:
        f.write(content)
    url = f"/static/uploads/ads/{filename}"
    return JSONResponse({"success": True, "url": url, "filename": filename})


@app.get("/blog/{post_slug}", response_class=HTMLResponse)
async def view_blog_post(post_slug: str, request: Request):
    """Public blog post view page"""
    with get_db() as db:
        post = db.execute("SELECT * FROM blog_posts WHERE slug=?", (post_slug,)).fetchone()
        if not post:
            from fastapi.responses import RedirectResponse
            return RedirectResponse("/dashboard")
        
        # Increment view count
        db.execute("UPDATE blog_posts SET view_count = view_count + 1 WHERE id=?", (post["id"],))
    
    user = get_current_user(request)
    if not user:
        from fastapi.responses import RedirectResponse
        return RedirectResponse("/login")
    
    from fastapi.templating import Jinja2Templates
    templates = Jinja2Templates(directory="/opt/exam_system/templates")
    # Get active ads for display
    with get_db() as db:
        ads = db.execute("SELECT id, title, image_url, link_url FROM ads WHERE is_active=1 ORDER BY sort_order ASC, id ASC").fetchall()
    
    return templates.TemplateResponse("blog_view.html", {"request": request, "post": dict(post), "user": user, "ads": [dict(a) for a in ads]})


# ============================================================
# Startup
# ============================================================


# ============================================================
# AI Question Generation
# ============================================================


def get_site_settings():
    with get_db() as db:
        row = db.execute("SELECT * FROM site_settings WHERE id=1").fetchone()
        if not row:
            db.execute("INSERT INTO site_settings (id) VALUES (1)")
            row = db.execute("SELECT * FROM site_settings WHERE id=1").fetchone()
        d = dict(row)
        # Also ensure columns exist (migration)
        cols = [r["name"] for r in db.execute("PRAGMA table_info(site_settings)").fetchall()]
        if "allow_register" not in cols:
            db.execute("ALTER TABLE site_settings ADD COLUMN allow_register INTEGER DEFAULT 1")
        if "allow_account_login" not in cols:
            db.execute("ALTER TABLE site_settings ADD COLUMN allow_account_login INTEGER DEFAULT 1")
        if "allow_student_id_login" not in cols:
            db.execute("ALTER TABLE site_settings ADD COLUMN allow_student_id_login INTEGER DEFAULT 1")
        user_cols2 = [c2['name'] for c2 in db.execute('PRAGMA table_info(users)').fetchall()]
        if 'plain_password' not in user_cols2:
            db.execute('ALTER TABLE users ADD COLUMN plain_password TEXT DEFAULT ""')
        if not row:
            row = db.execute("SELECT * FROM site_settings WHERE id=1").fetchone()
            d = dict(row)
        d["allow_register"] = bool(d.get("allow_register", 1))
        d["allow_account_login"] = bool(d.get("allow_account_login", 1))
        d["allow_student_id_login"] = bool(d.get("allow_student_id_login", 1))
        return d

def get_ai_settings():
    with get_db() as db:
        row = db.execute("SELECT * FROM ai_settings WHERE id=1").fetchone()
        return dict(row) if row else {"provider": "deepseek", "api_key": "", "model": "deepseek-chat", "base_url": "https://api.deepseek.com"}

@app.get("/api/admin/ai/settings")
async def get_ai_settings_api(request: Request):
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)
    settings = get_ai_settings()
    # Mask API key
    if settings["api_key"]:
        key = settings["api_key"]
        settings["api_key_masked"] = key[:6] + "****" + key[-4:] if len(key) > 10 else "****"
    else:
        settings["api_key_masked"] = ""
    return JSONResponse(settings)

@app.post("/api/admin/ai/settings")
async def save_ai_settings_api(request: Request):
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)
    data = await request.json()
    provider = data.get("provider", "deepseek")
    api_key = data.get("api_key", "")
    model = data.get("model", "")
    # Set default model based on provider
    if not model:
        if provider == "deepseek":
            model = "deepseek-chat"
        elif provider == "siliconflow":
            model = "deepseek-ai/DeepSeek-V3"
        else:
            model = "deepseek-chat"
    # Set base_url based on provider
    if provider == "deepseek":
        base_url = "https://api.deepseek.com"
    elif provider == "siliconflow":
        base_url = "https://api.siliconflow.cn/v1"
    else:
        base_url = data.get("base_url", "https://api.deepseek.com")

    with get_db() as db:
        db.execute("""UPDATE ai_settings SET provider=?, api_key=?, model=?, base_url=? WHERE id=1""",
                   (provider, api_key, model, base_url))
    return JSONResponse({"success": True})


@app.post("/api/admin/ai/textbook/upload")
async def upload_textbook(request: Request, file: UploadFile = File(...)):
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)

    filename = file.filename or "unknown"
    ext = filename.rsplit('.', 1)[-1].lower() if '.' in filename else ''

    if ext not in ('pdf', 'docx', 'doc', 'txt'):
        return JSONResponse({"error": "仅支持 PDF/Word/TXT 格式"}, status_code=400)

    content_bytes = await file.read()
    file_size = len(content_bytes)

    if file_size > 20 * 1024 * 1024:
        return JSONResponse({"error": "文件不能超过 20MB"}, status_code=400)

    text_content = ""
    try:
        if ext == 'pdf':
            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix='.pdf', delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name
            try:
                doc = fitz.open(tmp_path)
                pages_text = []
                for page in doc:
                    pages_text.append(page.get_text())
                text_content = "\n".join(pages_text)
                doc.close()
            finally:
                os.unlink(tmp_path)

        elif ext in ('docx', 'doc'):
            import tempfile, os
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
                tmp.write(content_bytes)
                tmp_path = tmp.name
            try:
                doc = DocxDocument(tmp_path)
                paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
                text_content = "\n".join(paragraphs)
            finally:
                os.unlink(tmp_path)

        elif ext == 'txt':
            text_content = content_bytes.decode('utf-8', errors='replace')

    except Exception as e:
        return JSONResponse({"error": f"文件解析失败: {str(e)}"}, status_code=500)

    if not text_content.strip():
        return JSONResponse({"error": "文件内容为空，无法提取文字"}, status_code=400)

    # Truncate if too long (AI context limit ~15000 chars)
    max_chars = 15000
    if len(text_content) > max_chars:
        text_content = text_content[:max_chars] + "\n...[内容过长，已截断]"

    char_count = len(text_content)

    with get_db() as db:
        cursor = db.execute(
            "INSERT INTO textbooks (file_name, file_size, content, char_count) VALUES (?, ?, ?, ?)",
            (filename, file_size, text_content, char_count)
        )
        tb_id = cursor.lastrowid

    return JSONResponse({"success": True, "id": tb_id, "file_name": filename, "char_count": char_count})


@app.get("/api/admin/ai/textbooks")
async def list_textbooks(request: Request):
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)

    with get_db() as db:
        rows = db.execute(
            "SELECT id, file_name, file_size, char_count, created_at FROM textbooks ORDER BY created_at DESC"
        ).fetchall()

    return JSONResponse({"textbooks": [dict(r) for r in rows]})


@app.delete("/api/admin/ai/textbooks/{textbook_id}")
async def delete_textbook(textbook_id: int, request: Request):
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)

    with get_db() as db:
        db.execute("DELETE FROM textbooks WHERE id=?", (textbook_id,))
    return JSONResponse({"success": True})

@app.post("/api/admin/ai/generate")
async def ai_generate_questions(request: Request):
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)

    settings = get_ai_settings()
    if not settings["api_key"]:
        return JSONResponse({"error": "请先在AI设置中配置API Key"}, status_code=400)

    data = await request.json()
    q_type = data.get("type", "single")
    q_types = [t.strip() for t in q_type.split(',') if t.strip()] if ',' in q_type else [q_type]
    category = data.get("category", "总论")
    topic = data.get("topic", "")
    textbook_ids = data.get("textbook_ids", [])
    if not textbook_ids:
        tb_id = data.get("textbook_id")
        if tb_id:
            textbook_ids = [tb_id]

    # Get textbook content from multiple textbooks
    textbook_content = ""
    if textbook_ids:
        with get_db() as db:
            for tid in textbook_ids:
                try:
                    tb = db.execute("SELECT content, file_name FROM textbooks WHERE id=?", (int(tid),)).fetchone()
                    if tb:
                        header = f"\n--- 《{tb['file_name']}》 ---\n" if len(textbook_ids) > 1 else ""
                        textbook_content += header + tb["content"] + "\n"
                except (ValueError, TypeError):
                    pass
    counts = {
        "single": min(int(data.get("count_single", 5)), 20),
        "multiple": min(int(data.get("count_multiple", 3)), 20),
        "shared": min(int(data.get("count_shared", 2)), 10),
        "case": min(int(data.get("count_case", 1)), 5),
    }
    difficulty = data.get("difficulty", "副高")
    level = data.get("level", "2")
    level_desc = {"1": "基础难度，侧重基本概念和常见知识点", "2": "中等难度，侧重综合分析和鉴别诊断", "3": "高难度，侧重疑难病例、前沿知识和综合判断"}
    level_hint = level_desc.get(str(level), level_desc["2"])

    type_names = {"single": "单选题(A-D四个选项)", "multiple": "多选题(A-H八个选项)", "shared": "共用题干单选题(含题干和3-5小题，每题A-H八个选项)", "case": "案例分析多选题(含病例题干和6-8小题，每题A-H八个选项)"}
    type_desc = "、".join([type_names.get(t, "单选题") for t in q_types])
    type_counts_desc = "、".join([f"{type_names.get(t, t)}{counts.get(t, 5)}{'组' if t in ('shared','case') else '题'}" for t in q_types])

    category_hint = "，涵盖病理学各系统分类（总论及各论）" if category == "套卷" else (f"分类为「{category}」" if category else "")
    topic_hint = f"，重点围绕「{topic}」" if topic else ""

    # Build textbook reference section
    textbook_section = ""
    if textbook_content:
        textbook_section = f"\n\n【参考教材内容】\n{textbook_content[:8000]}\n【教材内容结束】\n"

    # Build textbook name list for non-streaming endpoint
    textbook_names = ""
    if textbook_ids:
        tb_names = []
        with get_db() as db:
            for tid in textbook_ids:
                try:
                    tb2 = db.execute("SELECT file_name FROM textbooks WHERE id=?", (int(tid),)).fetchone()
                    if tb2:
                        tb_names.append(tb2["file_name"])
                except (ValueError, TypeError):
                    pass
        if tb_names:
            textbook_names = "、".join(tb_names)

    # Build structured type+count list
    type_count_lines = []
    for t in q_types:
        c = counts.get(t, 5)
        name = type_names.get(t, t)
        unit = "组" if t in ("shared", "case") else "题"
        type_count_lines.append("  - " + name + ": " + str(c) + unit)
    type_count_text = "\n".join(type_count_lines)

    level_labels2 = {"1": "1级·基础", "2": "2级·中等", "3": "3级·困难"}
    level_label = level_labels2.get(str(level), "2级·中等")

    prompt = "【出题任务】你是一位资深病理学教授，请严格按照以下要求生成病理学考试题。\n\n"
    prompt += "【考试级别】" + difficulty + "\n"
    prompt += "【难度级别】" + level_label + " — " + level_hint + "\n"
    prompt += "【题目分类】" + (category if category else "不限") + ("（套卷：涵盖病理学所有分类）" if category == "套卷" else "") + "\n"
    if topic:
        prompt += "【出题主题】重点围绕「" + topic + "」\n"
    prompt += "\n【题型与数量】\n" + type_count_text + "\n"
    if textbook_names:
        prompt += "\n【参考教材】" + textbook_names + "\n"
    prompt += "\n【出题要求】\n"
    prompt += "1. 题目内容准确、专业，符合病理学" + difficulty + "级别、" + level_label + "难度的考试要求\n"
    prompt += "2. " + ("必须严格基于上述参考教材内容出题，题目知识点要覆盖教材中的重点" if textbook_names else "基于病理学专业知识出题") + "\n"
    prompt += "3. 每道题必须有明确的正确答案和简要解析\n"
    prompt += "4. 选项要有干扰性，体现考试出题水平\n"
    prompt += "5. 严格按照以下JSON格式返回，不要返回任何其他内容\n"
    if textbook_section:
        prompt += "\n【教材内容】\n" + textbook_section + "\n"
    prompt += "\n【JSON返回格式】\n"
    prompt += '单选题和多选题格式：{"question":"题目", "options":{"A":"选项A",...}, "answer":"A", "explanation":"解析", "knowledge_point":"知识点"}\n'
    prompt += '共用题干和案例分析格式：{"stem":"题干", "sub_questions":[{"question":"小题", "options":{...}, "answer":"A", "explanation":"解析"}]}\n'
    prompt += "所有题目放在一个JSON数组中返回。只返回JSON，不要有其他文字。"

    try:
        async with httpx.AsyncClient(timeout=120.0) as client:
            resp = await client.post(
                f"{settings['base_url']}/chat/completions",
                headers={
                    "Authorization": f"Bearer {settings['api_key']}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": settings["model"],
                    "messages": [
                        {"role": "system", "content": "你是一位资深的病理学教授，擅长根据教材内容出病理学高级职称考试题。请严格基于提供的教材内容出题，按照要求的JSON格式返回。"},
                        {"role": "user", "content": prompt}
                    ],
                    "temperature": 0.7,
                    "max_tokens": 6000
                }
            )
            resp.raise_for_status()
            result = resp.json()
            ai_text = result["choices"][0]["message"]["content"]

            # Extract JSON from response
            import json as json_lib
            # Try to find JSON array in the response
            json_match = re.search(r'\[.*\]', ai_text, re.DOTALL)
            if json_match:
                questions = json_lib.loads(json_match.group())
                return JSONResponse({"success": True, "questions": questions, "type": q_type})
            else:
                return JSONResponse({"error": "AI返回格式异常，请重试", "raw": ai_text[:500]}, status_code=500)

    except httpx.TimeoutException:
        return JSONResponse({"error": "AI请求超时，请稍后重试"}, status_code=504)
    except httpx.HTTPStatusError as e:
        if e.response.status_code == 401:
            return JSONResponse({"error": "API Key无效，请检查设置"}, status_code=401)
        return JSONResponse({"error": f"AI服务错误: {e.response.status_code}"}, status_code=500)
    except Exception as e:
        return JSONResponse({"error": f"生成失败: {str(e)}"}, status_code=500)



@app.post("/api/admin/ai/generate-stream")
async def ai_generate_stream(request: Request):
    """Stream AI generation output via SSE"""
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)

    settings = get_ai_settings()
    if not settings["api_key"]:
        return JSONResponse({"error": "请先在AI设置中配置API Key"}, status_code=400)

    data = await request.json()
    q_type = data.get("type", "single")
    q_types = [t.strip() for t in q_type.split(',') if t.strip()] if ',' in q_type else [q_type]
    category = data.get("category", "总论")
    topic = data.get("topic", "")
    textbook_ids = data.get("textbook_ids", [])
    if not textbook_ids:
        tb_id = data.get("textbook_id")
        if tb_id:
            textbook_ids = [tb_id]

    textbook_content = ""
    if textbook_ids:
        with get_db() as db:
            for tid in textbook_ids:
                try:
                    tb = db.execute("SELECT content, file_name FROM textbooks WHERE id=?", (int(tid),)).fetchone()
                    if tb:
                        header = "\n--- 《" + tb["file_name"] + "》 ---\n" if len(textbook_ids) > 1 else ""
                        textbook_content += header + tb["content"] + "\n"
                except (ValueError, TypeError):
                    pass

    counts = {
        "single": min(int(data.get("count_single", 5)), 20),
        "multiple": min(int(data.get("count_multiple", 3)), 20),
        "shared": min(int(data.get("count_shared", 2)), 10),
        "case": min(int(data.get("count_case", 1)), 5),
    }
    difficulty = data.get("difficulty", "副高")
    level = data.get("level", "2")
    level_desc = {"1": "基础难度，侧重基本概念和常见知识点", "2": "中等难度，侧重综合分析和鉴别诊断", "3": "高难度，侧重疑难病例、前沿知识和综合判断"}
    level_hint = level_desc.get(str(level), level_desc["2"])

    type_names = {"single": "单选题(A-D四个选项)", "multiple": "多选题(A-H八个选项)", "shared": "共用题干单选题(含题干和3-5小题，每题A-H八个选项)", "case": "案例分析多选题(含病例题干和6-8小题，每题A-H八个选项)"}
    type_counts_desc = "、".join([type_names.get(t, t) + str(counts.get(t, 5)) + ("组" if t in ("shared","case") else "题") for t in q_types])

    category_hint = "，涵盖病理学各系统分类（总论及各论）" if category == "套卷" else ("分类为「" + category + "」" if category else "")
    topic_hint = "，重点围绕「" + topic + "」" if topic else ""

    textbook_section = ""
    if textbook_content:
        textbook_section = "\n\n【参考教材内容】\n" + textbook_content[:8000] + "\n【教材内容结束】\n"

    # Build textbook name list
    textbook_names = ""
    if textbook_ids:
        tb_names = []
        with get_db() as db:
            for tid in textbook_ids:
                try:
                    tb = db.execute("SELECT file_name FROM textbooks WHERE id=?", (int(tid),)).fetchone()
                    if tb:
                        tb_names.append(tb["file_name"])
                except (ValueError, TypeError):
                    pass
        if tb_names:
            textbook_names = "、".join(tb_names)

    # Build structured type+count list
    type_count_lines = []
    for t in q_types:
        c = counts.get(t, 5)
        name = type_names.get(t, t)
        unit = "组" if t in ("shared", "case") else "题"
        type_count_lines.append("  - " + name + ": " + str(c) + unit)
    type_count_text = "\n".join(type_count_lines)

    level_labels = {"1": "1级·基础", "2": "2级·中等", "3": "3级·困难"}
    level_label = level_labels.get(str(level), "2级·中等")

    prompt = "【出题任务】你是一位资深病理学教授，请严格按照以下要求生成病理学考试题。\n\n"
    prompt += "【考试级别】" + difficulty + "\n"
    prompt += "【难度级别】" + level_label + " — " + level_hint + "\n"
    prompt += "【题目分类】" + (category if category else "不限") + ("（套卷：涵盖病理学所有分类）" if category == "套卷" else "") + "\n"
    if topic:
        prompt += "【出题主题】重点围绕「" + topic + "」\n"
    prompt += "\n【题型与数量】\n" + type_count_text + "\n"
    if textbook_names:
        prompt += "\n【参考教材】" + textbook_names + "\n"
    prompt += "\n【出题要求】\n"
    prompt += "1. 题目内容准确、专业，符合病理学" + difficulty + "级别、" + level_label + "难度的考试要求\n"
    prompt += "2. " + ("必须严格基于上述参考教材内容出题，题目知识点要覆盖教材中的重点" if textbook_names else "基于病理学专业知识出题") + "\n"
    prompt += "3. 每道题必须有明确的正确答案和简要解析\n"
    prompt += "4. 选项要有干扰性，体现考试出题水平\n"
    prompt += "5. 严格按照以下JSON格式返回，不要返回任何其他内容\n"
    if textbook_section:
        prompt += "\n【教材内容】\n" + textbook_section + "\n"
    prompt += "\n【JSON返回格式】\n"
    prompt += "单选题和多选题：{\"question\":\"题目\", \"options\":{\"A\":\"选项A\",...}, \"answer\":\"A\", \"explanation\":\"解析\", \"knowledge_point\":\"知识点\"}\n"
    prompt += "共用题干和案例分析：{\"stem\":\"题干\", \"sub_questions\":[{\"question\":\"小题\", \"options\":{...}, \"answer\":\"A\", \"explanation\":\"解析\"}]}\n"
    prompt += "所有题目放在一个JSON数组中返回。只返回JSON，不要有其他文字。"

    async def event_generator():
        full_text = ""
        try:
            async with httpx.AsyncClient(timeout=180.0) as client:
                async with client.stream(
                    "POST",
                    settings["base_url"] + "/chat/completions",
                    headers={
                        "Authorization": "Bearer " + settings["api_key"],
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": settings["model"],
                        "messages": [
                            {"role": "system", "content": "你是一位资深的病理学教授，擅长根据教材内容出病理学高级职称考试题。请严格基于提供的教材内容出题，按照要求的JSON格式返回。"},
                            {"role": "user", "content": prompt}
                        ],
                        "temperature": 0.7,
                        "max_tokens": 8000,
                        "stream": True
                    }
                ) as response:
                    if response.status_code != 200:
                        error_body = await response.aread()
                        err_msg = "API错误: " + str(response.status_code) + " " + error_body.decode()[:200]
                        yield "data: " + json.dumps({"type": "error", "error": err_msg}) + "\n\n"
                        return

                    async for line in response.aiter_lines():
                        if not line.startswith("data: "):
                            continue
                        payload = line[6:]
                        if payload.strip() == "[DONE]":
                            break
                        try:
                            chunk = json.loads(payload)
                            delta = chunk.get("choices", [{}])[0].get("delta", {})
                            content = delta.get("content", "")
                            if content:
                                full_text += content
                                yield "data: " + json.dumps({"type": "chunk", "content": content}, ensure_ascii=False) + "\n\n"
                        except (json.JSONDecodeError, IndexError, KeyError):
                            continue

            # Parse the full JSON response
            json_match = re.search(r'\[.*\]', full_text, re.DOTALL)
            if json_match:
                questions = json.loads(json_match.group())
                yield "data: " + json.dumps({"type": "done", "questions": questions, "type_name": q_type}, ensure_ascii=False) + "\n\n"
            else:
                yield "data: " + json.dumps({"type": "error", "error": "AI返回格式异常，无法解析JSON", "raw": full_text[:500]}, ensure_ascii=False) + "\n\n"

        except httpx.TimeoutException:
            yield "data: " + json.dumps({"type": "error", "error": "AI请求超时，请稍后重试"}) + "\n\n"
        except Exception as e:
            yield "data: " + json.dumps({"type": "error", "error": "生成异常: " + str(e)}, ensure_ascii=False) + "\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream",
                             headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@app.post("/api/admin/ai/import")
async def ai_import_questions(request: Request):
    """Import AI-generated questions into the question bank"""
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)

    data = await request.json()
    questions = data.get("questions", [])
    q_type = data.get("type", "single")
    category = data.get("category", "总论")

    if not questions:
        return JSONResponse({"error": "没有题目可导入"}, status_code=400)

    imported = 0
    errors = []

    with get_db() as db:
        for i, q in enumerate(questions):
            try:
                if q_type in ("single", "multiple"):
                    question_text = q.get("question", "")
                    options = q.get("options", {})
                    answer = q.get("answer", "")
                    explanation = q.get("explanation", "")

                    if not question_text or not options or not answer:
                        errors.append(f"第{i+1}题：缺少必要字段")
                        continue

                    # Format options as A. xxx\nB. xxx
                    opt_lines = []
                    for k in sorted(options.keys()):
                        opt_lines.append(f"{k}. {options[k]}")
                    options_text = "\n".join(opt_lines)

                    db.execute("""INSERT INTO questions (type, category, question, options, answer, explanation)
                                  VALUES (?, ?, ?, ?, ?, ?)""",
                               (q_type, category, question_text, options_text, answer, explanation))
                    imported += 1

                elif q_type in ("shared", "case"):
                    stem = q.get("stem", "")
                    subs = q.get("sub_questions", [])
                    if not stem or not subs:
                        errors.append(f"第{i+1}组：缺少题干或小题")
                        continue

                    # Create a group question
                    all_questions = []
                    all_answers = []
                    all_explanations = []

                    for sq in subs:
                        sq_text = f"[共用题干] {stem}\n{sq.get('question', '')}"
                        options = sq.get("options", {})
                        opt_lines = []
                        for k in sorted(options.keys()):
                            opt_lines.append(f"{k}. {options[k]}")
                        options_text = "\n".join(opt_lines)
                        answer = sq.get("answer", "")
                        explanation = sq.get("explanation", "")

                        db.execute("""INSERT INTO questions (type, category, question, options, answer, explanation)
                                      VALUES (?, ?, ?, ?, ?, ?)""",
                                   (q_type, category, sq_text, options_text, answer, explanation))
                        imported += 1

            except Exception as e:
                errors.append(f"第{i+1}题导入失败: {str(e)}")

    return JSONResponse({"success": True, "imported": imported, "errors": errors})



@app.post("/api/admin/ai/import-as-exam")
async def ai_import_as_exam(request: Request):
    """Import AI-generated questions as a complete exam paper (for 套卷)"""
    user = get_current_user(request)
    if not user or user["role"] != "admin":
        return JSONResponse({"error": "权限不足"}, status_code=403)

    data = await request.json()
    questions = data.get("questions", [])
    q_type = data.get("type", "single")
    category = data.get("category", "套卷")
    title = data.get("title", "").strip()
    difficulty = data.get("difficulty", "副高")
    level = data.get("level", "2")

    if not questions:
        return JSONResponse({"error": "没有题目可导入"}, status_code=400)

    if not title:
        level_labels = {"1": "基础", "2": "中等", "3": "困难"}
        ll = level_labels.get(str(level), "中等")
        title = difficulty + "职称模拟试卷（" + ll + "难度）"

    imported = 0
    errors = []
    question_ids = []

    with get_db() as db:
        # Step 1: Import all questions into question bank
        for i, q in enumerate(questions):
            try:
                # Determine the actual type of this question
                actual_type = q_type
                if q.get("stem"):
                    actual_type = "shared" if q_type in ("single", "multiple", "shared") else "case"
                elif q.get("question"):
                    actual_type = "single" if q_type in ("single",) else q_type

                if actual_type in ("single", "multiple") or (not q.get("stem") and q.get("question")):
                    question_text = q.get("question", "")
                    options = q.get("options", {})
                    answer = q.get("answer", "")
                    explanation = q.get("explanation", "")

                    if not question_text or not options or not answer:
                        errors.append("第" + str(i+1) + "题：缺少必要字段")
                        continue

                    opt_lines = []
                    for k in sorted(options.keys()):
                        opt_lines.append(k + ". " + options[k])
                    options_text = "\n".join(opt_lines)

                    # Determine actual type from answer format
                    q_actual_type = "single"
                    if len(answer) > 1:
                        q_actual_type = "multiple"

                    cursor = db.execute(
                        "INSERT INTO questions (type, category, question, options, answer, explanation) VALUES (?, ?, ?, ?, ?, ?)",
                        (q_actual_type, category, question_text, options_text, answer, explanation))
                    question_ids.append(cursor.lastrowid)
                    imported += 1

                elif q.get("stem") and q.get("sub_questions"):
                    stem = q["stem"]
                    subs = q["sub_questions"]
                    actual_type = "shared" if actual_type in ("single", "shared") else "case"

                    for sq in subs:
                        sq_text = "[共用题干] " + stem + "\n" + sq.get("question", "")
                        options = sq.get("options", {})
                        opt_lines = []
                        for k in sorted(options.keys()):
                            opt_lines.append(k + ". " + options[k])
                        options_text = "\n".join(opt_lines)
                        answer = sq.get("answer", "")
                        explanation = sq.get("explanation", "")

                        cursor = db.execute(
                            "INSERT INTO questions (type, category, question, options, answer, explanation) VALUES (?, ?, ?, ?, ?, ?)",
                            (actual_type, category, sq_text, options_text, answer, explanation))
                        question_ids.append(cursor.lastrowid)
                        imported += 1

            except Exception as e:
                errors.append("第" + str(i+1) + "题导入失败: " + str(e))

        if not question_ids:
            return JSONResponse({"error": "没有成功导入任何题目", "errors": errors}, status_code=400)

        # Step 2: Create exam
        duration = max(60, len(question_ids) * 2)
        total_score = float(len(question_ids))
        pass_score = round(total_score * 0.6, 1)

        cursor = db.execute(
            "INSERT INTO exams (title, description, duration, total_score, pass_score, is_practice, shuffle_questions, status) VALUES (?, ?, ?, ?, ?, ?, ?, ?)",
            (title, "AI生成套卷 - " + difficulty + " " + category, duration, total_score, pass_score, False, True, "draft"))
        exam_id = cursor.lastrowid

        # Step 3: Link questions to exam
        for idx, qid in enumerate(question_ids):
            db.execute(
                "INSERT INTO exam_questions (exam_id, question_id, display_order, score_override) VALUES (?, ?, ?, ?)",
                (exam_id, qid, idx + 1, 1.0))

    return JSONResponse({
        "success": True,
        "imported": imported,
        "exam_id": exam_id,
        "exam_title": title,
        "question_count": len(question_ids),
        "errors": errors
    })




@app.get("/api/admin/exams/{exam_id}/export")
async def export_exam(exam_id: int, request: Request, mode: str = "all"):
    """Export exam as Word document. mode: 'exam' = questions only, 'answer' = with answers, 'all' = both"""
    user = require_admin(request)
    from docx import Document as DocxDocument
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    with get_db() as db:
        exam = db.execute("SELECT * FROM exams WHERE id=?", (exam_id,)).fetchone()
        if not exam:
            return JSONResponse({"error": "试卷不存在"}, status_code=404)
        eqs = db.execute("SELECT q.*, eq.score_override FROM exam_questions eq JOIN questions q ON eq.question_id=q.id WHERE eq.exam_id=? ORDER BY eq.display_order", (exam_id,)).fetchall()

    if not eqs:
        return JSONResponse({"error": "试卷无题目"}, status_code=400)

    def build_doc(include_answer):
        doc = DocxDocument()
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.font.size = Pt(12)
        style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(exam["title"])
        run.font.size = Pt(18)
        run.bold = True
        run.font.name = '宋体'
        run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_text = "满分：" + str(exam['total_score']) + "分    时长：" + str(exam['duration']) + "分钟    合格线：" + str(exam['pass_score']) + "分"
        run = sub.add_run(info_text)
        run.font.size = Pt(11)
        run.font.name = '宋体'
        run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

        doc.add_paragraph()

        type_names = {"single": "一、单选题", "multiple": "二、多选题", "shared": "三、共用题干单选题", "case": "四、案例分析题"}
        type_order = ["single", "multiple", "shared", "case"]
        grouped = {}
        for q in eqs:
            t = q["type"]
            if t not in grouped:
                grouped[t] = []
            grouped[t].append(q)

        q_num = 1
        for t in type_order:
            if t not in grouped:
                continue
            questions = grouped[t]
            h = doc.add_paragraph()
            run = h.add_run(type_names.get(t, t))
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

            for q in questions:
                q_text = q["question"].replace("[共用题干] ", "")
                p = doc.add_paragraph()
                run = p.add_run(str(q_num) + ". " + q_text)
                run.font.size = Pt(12)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

                if q["options"]:
                    for line in q["options"].split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        op = doc.add_paragraph()
                        run = op.add_run("    " + line)
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
                        run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

                if include_answer:
                    ans_p = doc.add_paragraph()
                    ans_text = "【答案】" + q['answer']
                    if q["explanation"]:
                        ans_text += "    【解析】" + q["explanation"]
                    run = ans_p.add_run(ans_text)
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(0, 0, 180)
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

                q_num += 1
            doc.add_paragraph()

        return doc

    if mode == "exam":
        doc = build_doc(include_answer=False)
        filename = exam["title"] + "_试题卷.docx"
    elif mode == "answer":
        doc = build_doc(include_answer=True)
        filename = exam["title"] + "_答案解析卷.docx"
    else:
        doc = build_doc(include_answer=False)
        doc.add_page_break()
        title2 = doc.add_paragraph()
        title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title2.add_run("答案与解析")
        run.font.size = Pt(18)
        run.bold = True
        run.font.name = '宋体'
        run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
        doc.add_paragraph()
        type_names = {"single": "一、单选题", "multiple": "二、多选题", "shared": "三、共用题干单选题", "case": "四、案例分析题"}
        type_order = ["single", "multiple", "shared", "case"]
        grouped = {}
        for q in eqs:
            t = q["type"]
            if t not in grouped:
                grouped[t] = []
            grouped[t].append(q)
        q_num = 1
        for t in type_order:
            if t not in grouped:
                continue
            questions = grouped[t]
            h = doc.add_paragraph()
            run = h.add_run(type_names.get(t, t))
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = '宋体'
            run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
            for q in questions:
                q_text = q["question"].replace("[共用题干] ", "")
                p = doc.add_paragraph()
                run = p.add_run(str(q_num) + ". " + q_text)
                run.font.size = Pt(12)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
                if q["options"]:
                    for line in q["options"].split("\n"):
                        line = line.strip()
                        if not line:
                            continue
                        op = doc.add_paragraph()
                        run = op.add_run("    " + line)
                        run.font.size = Pt(12)
                        run.font.name = '宋体'
                        run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
                ans_p = doc.add_paragraph()
                ans_text = "【答案】" + q['answer']
                if q["explanation"]:
                    ans_text += "    【解析】" + q["explanation"]
                run = ans_p.add_run(ans_text)
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0, 0, 180)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')
                q_num += 1
            doc.add_paragraph()
        filename = exam["title"] + "_完整卷.docx"

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    from fastapi.responses import StreamingResponse
    import urllib.parse
    encoded_filename = urllib.parse.quote(filename)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename*=UTF-8''" + encoded_filename}
    )


@app.on_event("startup")
async def startup():
    init_db()

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8090)
